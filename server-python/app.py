"""
PDF转Word服务端 - 增强版
使用pdf2docx库实现高质量PDF转换
支持多种转换模式、页码选择、PDF预览等功能
"""
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from pdf2docx import Converter
import fitz  # PyMuPDF
import os
import uuid
import logging
import io
import base64
from werkzeug.utils import secure_filename
from datetime import datetime
from PIL import Image, ImageDraw, ImageFont
import time
import traceback
import re
import shutil
import platform
try:
    import requests
    HAS_REQUESTS = True
except ImportError:
    HAS_REQUESTS = False

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app)  # 允许跨域请求

# 配置
UPLOAD_FOLDER = 'temp/uploads'
CONVERTED_FOLDER = 'temp/converted'
ALLOWED_EXTENSIONS = {'pdf'}
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB

# 确保目录存在
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(CONVERTED_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['CONVERTED_FOLDER'] = CONVERTED_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_FILE_SIZE


def allowed_file(filename):
    """检查文件扩展名是否允许"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def cleanup_old_files(folder, max_age_hours=24):
    """清理超过指定时间的临时文件"""
    try:
        now = datetime.now().timestamp()
        for filename in os.listdir(folder):
            file_path = os.path.join(folder, filename)
            if os.path.isfile(file_path):
                file_age = now - os.path.getmtime(file_path)
                if file_age > max_age_hours * 3600:
                    os.remove(file_path)
                    logger.info(f"已清理旧文件: {filename}")
    except Exception as e:
        logger.error(f"清理文件失败: {str(e)}")


@app.route('/health', methods=['GET'])
def health_check():
    """健康检查端点"""
    return jsonify({
        'status': 'UP',
        'service': 'PDF to Word Converter (pdf2docx) - Enhanced + High Quality + Fast',
        'version': '2.2.0',
        'features': ['complex_format', 'images', 'page_selection', 'preview', 'premium_mode', 'fast_mode', 'multi_processing'],
        'modes': {
            'premium': '⭐高质量模式（表格+图片+样式，质量与速度平衡，推荐！）',
            'complex': '完整格式（表格+图片+样式，传统模式）',
            'fast': '快速模式（表格优化，速度优先）',
            'simple': '简化格式（表格+样式，无图片）',
            'text-only': '纯文本（仅文字）'
        },
        'recommended': 'premium'
    })


@app.route('/pdf/info', methods=['POST'])
def get_pdf_info():
    """
    获取PDF信息和预览图
    返回页数、每页缩略图等信息
    """
    try:
        if 'file' not in request.files:
            return jsonify({'error': '没有上传文件'}), 400
        
        file = request.files['file']
        
        if file.filename == '' or not allowed_file(file.filename):
            return jsonify({'error': '无效的PDF文件'}), 400
        
        # 临时保存文件
        file_uuid = str(uuid.uuid4())
        pdf_filename = f"{file_uuid}_temp.pdf"
        pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], pdf_filename)
        file.save(pdf_path)
        
        try:
            # 打开PDF获取信息
            doc = fitz.open(pdf_path)
            page_count = len(doc)
            
            # 生成每页的缩略图（base64编码）
            previews = []
            for page_num in range(min(page_count, 20)):  # 最多返回前20页预览
                page = doc[page_num]
                
                # 生成缩略图 (150x200 像素)
                zoom = 150 / page.rect.width
                mat = fitz.Matrix(zoom, zoom)
                pix = page.get_pixmap(matrix=mat, alpha=False)
                
                # 转换为PIL Image并压缩
                img_data = pix.tobytes("png")
                img = Image.open(io.BytesIO(img_data))
                
                # 进一步压缩
                buffer = io.BytesIO()
                img.save(buffer, format="JPEG", quality=70, optimize=True)
                img_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
                
                previews.append({
                    'page': page_num + 1,
                    'image': f'data:image/jpeg;base64,{img_base64}',
                    'width': int(page.rect.width),
                    'height': int(page.rect.height)
                })
            
            doc.close()
            
            # 删除临时文件
            os.remove(pdf_path)
            
            logger.info(f"PDF信息获取成功: {page_count}页")
            
            return jsonify({
                'success': True,
                'pageCount': page_count,
                'previews': previews,
                'filename': file.filename
            })
            
        except Exception as e:
            logger.error(f"PDF信息获取失败: {str(e)}")
            try:
                os.remove(pdf_path)
            except:
                pass
            return jsonify({'error': f'PDF处理失败: {str(e)}'}), 500
            
    except Exception as e:
        logger.error(f"请求处理失败: {str(e)}")
        return jsonify({'error': f'请求失败: {str(e)}'}), 500


@app.route('/pdf/preview', methods=['POST'])
def get_pdf_preview_single_page():
    """
    获取PDF单页预览（按需加载，性能优化）
    只返回指定页的预览图
    """
    try:
        if 'file' not in request.files:
            return jsonify({'error': '没有上传文件'}), 400
        
        file = request.files['file']
        page_num = request.form.get('page', '1')
        
        try:
            page_num = int(page_num)
        except:
            return jsonify({'error': '页码格式错误'}), 400
        
        if file.filename == '' or not allowed_file(file.filename):
            return jsonify({'error': '无效的PDF文件'}), 400
        
        # 临时保存文件
        file_uuid = str(uuid.uuid4())
        pdf_filename = f"{file_uuid}_temp.pdf"
        pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], pdf_filename)
        file.save(pdf_path)
        
        try:
            # 打开PDF
            doc = fitz.open(pdf_path)
            page_count = len(doc)
            
            # 检查页码有效性
            if page_num < 1 or page_num > page_count:
                doc.close()
                os.remove(pdf_path)
                return jsonify({'error': f'页码超出范围 (1-{page_count})'}), 400
            
            # 生成指定页的预览
            page = doc[page_num - 1]  # 索引从0开始
            
            # 生成较大的预览图 (200x280像素)
            zoom = 200 / page.rect.width
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            
            # 转换为PIL Image并压缩
            img_data = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_data))
            
            # 压缩为JPEG
            buffer = io.BytesIO()
            img.save(buffer, format="JPEG", quality=75, optimize=True)
            img_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
            
            preview = {
                'page': page_num,
                'image': f'data:image/jpeg;base64,{img_base64}',
                'width': int(page.rect.width),
                'height': int(page.rect.height)
            }
            
            doc.close()
            os.remove(pdf_path)
            
            logger.info(f"单页预览生成成功: 第{page_num}页")
            
            return jsonify({
                'success': True,
                'preview': preview,
                'pageCount': page_count
            })
            
        except Exception as e:
            logger.error(f"PDF预览生成失败: {str(e)}")
            try:
                os.remove(pdf_path)
            except:
                pass
            return jsonify({'error': f'预览生成失败: {str(e)}'}), 500
            
    except Exception as e:
        logger.error(f"请求处理失败: {str(e)}")
        return jsonify({'error': f'请求失败: {str(e)}'}), 500


@app.route('/pdf/toword', methods=['POST'])
def convert_pdf_to_word():
    """
    PDF转Word端点 - 增强版（性能优化 v2.1）
    支持多种转换模式和页码选择，针对表格文档进行性能优化
    
    参数:
        file: PDF文件 (multipart/form-data)
        mode: 转换模式 (premium/complex/fast/simple/text-only)
            - premium: 高质量模式（表格+图片，质量优先，推荐！）⭐
            - complex: 完整格式（保留表格、图片、格式）
            - fast: 快速模式（表格优化，速度优先）
            - simple: 简化格式（基本格式，不含图片）
            - text-only: 纯文本模式（仅提取文字）
        pages: 页码范围 (可选)
            - "all": 全部页（默认）
            - "1,3,5": 指定页码
            - "1-5": 页码范围
            - "1-3,5,7-9": 混合格式
        include_images: 是否包含图片 (true/false)
        
    返回:
        {url, filename, size, conversion_time, mode, pages_converted}
    """
    try:
        # 检查是否有文件
        if 'file' not in request.files:
            return jsonify({'error': '没有上传文件'}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({'error': '文件名为空'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'error': '只支持PDF文件'}), 400
        
        # 获取转换参数
        mode = request.form.get('mode', 'complex')  # complex, simple, text-only
        pages_param = request.form.get('pages', 'all')
        include_images = request.form.get('include_images', 'true').lower() == 'true'
        
        # 生成唯一文件名
        original_filename = secure_filename(file.filename)
        file_uuid = str(uuid.uuid4())
        pdf_filename = f"{file_uuid}_{original_filename}"
        word_filename = f"{file_uuid}_{os.path.splitext(original_filename)[0]}.docx"
        
        pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], pdf_filename)
        word_path = os.path.join(app.config['CONVERTED_FOLDER'], word_filename)
        
        # 保存上传的PDF文件
        file.save(pdf_path)
        file_size = os.path.getsize(pdf_path) / 1024
        logger.info(f"接收到PDF文件: {original_filename} ({file_size:.2f} KB)")
        logger.info(f"转换模式: {mode}, 页码: {pages_param}, 包含图片: {include_images}")
        
        # 清理旧文件
        cleanup_old_files(app.config['UPLOAD_FOLDER'])
        cleanup_old_files(app.config['CONVERTED_FOLDER'])
        
        # 解析页码范围
        start_page, end_page, selected_pages = parse_page_range(pdf_path, pages_param)
        
        # 开始转换
        logger.info(f"开始转换: {original_filename} -> {word_filename}")
        start_time = datetime.now()
        
        try:
            # 根据模式选择转换方式
            if mode == 'text-only':
                # 纯文本模式：使用PyMuPDF直接提取文本
                convert_pdf_to_text(pdf_path, word_path, start_page, end_page)
                
            elif mode == 'premium':
                # 高质量模式：平衡质量和速度，适合大多数文档
                logger.info("使用高质量模式（Premium）- 推荐模式")
                cv = Converter(pdf_path)
                cv.convert(
                    word_path, 
                    start=start_page, 
                    end=end_page, 
                    image=True,  # 包含图片（高质量）
                    multi_processing=True,  # 启用多进程
                    cpu_count=3,  # 使用3个CPU核心（性能优化）
                )
                cv.close()
                logger.info("高质量模式转换完成")
                
            elif mode == 'fast':
                # 快速模式：专为表格密集型文档优化（速度优先）
                logger.info("使用快速模式（Fast）- 表格优化")
                cv = Converter(pdf_path)
                cv.convert(
                    word_path, 
                    start=start_page, 
                    end=end_page, 
                    image=False,  # 快速模式不包含图片以提升速度
                    multi_processing=True,  # 启用多进程
                    cpu_count=3,  # 使用3个CPU核心（更激进的并行）
                )
                cv.close()
                logger.info("快速模式转换完成")
                
            elif mode == 'simple':
                # 简化模式：不包含图片，启用性能优化
                cv = Converter(pdf_path)
                cv.convert(
                    word_path, 
                    start=start_page, 
                    end=end_page, 
                    image=False,  # 不包含图片
                    multi_processing=True,  # 启用多进程加速
                    cpu_count=2  # 使用2个CPU核心
                )
                cv.close()
                
            else:  # complex mode (default)
                # 复杂模式：完整转换，启用性能优化
                cv = Converter(pdf_path)
                cv.convert(
                    word_path, 
                    start=start_page, 
                    end=end_page,
                    image=include_images,  # 根据参数决定是否包含图片
                    multi_processing=True,  # 启用多进程加速
                    cpu_count=2  # 使用2个CPU核心（平衡速度和资源）
                )
                cv.close()
            
            conversion_time = (datetime.now() - start_time).total_seconds()
            word_size = os.path.getsize(word_path)
            
            pages_converted = f"{start_page + 1}-{end_page if end_page else 'end'}" if pages_param != 'all' else 'all'
            
            logger.info(f"转换成功: {word_filename} ({word_size / 1024:.2f} KB, 耗时 {conversion_time:.2f}s)")
            
            # 删除临时PDF文件
            try:
                os.remove(pdf_path)
            except:
                pass
            
            # 返回下载链接
            return jsonify({
                'url': f'/download/{word_filename}',
                'filename': word_filename,
                'size': word_size,
                'conversion_time': f'{conversion_time:.2f}s',
                'mode': mode,
                'pages_converted': pages_converted,
                'include_images': include_images
            })
            
        except Exception as e:
            logger.error(f"转换失败: {str(e)}")
            # 清理文件
            try:
                os.remove(pdf_path)
            except:
                pass
            return jsonify({'error': f'转换失败: {str(e)}'}), 500
        
    except Exception as e:
        logger.error(f"处理请求失败: {str(e)}")
        return jsonify({'error': f'处理请求失败: {str(e)}'}), 500


def parse_page_range(pdf_path, pages_param):
    """
    解析页码参数
    返回 (start_page, end_page, selected_pages)
    """
    try:
        # 获取PDF总页数
        doc = fitz.open(pdf_path)
        total_pages = len(doc)
        doc.close()
        
        if pages_param == 'all' or not pages_param:
            return 0, None, list(range(total_pages))
        
        # 解析页码范围
        # 支持格式: "1,3,5" 或 "1-5" 或 "1-3,5,7-9"
        selected_pages = []
        for part in pages_param.split(','):
            part = part.strip()
            if '-' in part:
                # 范围
                start, end = part.split('-')
                start = int(start) - 1  # 转换为0-based索引
                end = int(end)
                selected_pages.extend(range(max(0, start), min(total_pages, end)))
            else:
                # 单个页码
                page = int(part) - 1  # 转换为0-based索引
                if 0 <= page < total_pages:
                    selected_pages.append(page)
        
        if not selected_pages:
            return 0, None, list(range(total_pages))
        
        # pdf2docx使用start和end参数（end可以为None表示到末尾）
        start_page = min(selected_pages)
        end_page = max(selected_pages) + 1 if selected_pages else None
        
        return start_page, end_page, selected_pages
        
    except Exception as e:
        logger.error(f"页码解析失败: {str(e)}")
        return 0, None, []


def convert_pdf_to_text(pdf_path, output_path, start_page=0, end_page=None):
    """
    纯文本模式转换：提取PDF文本并保存为Word文档
    优化版：每页PDF内容保留在Word的同一页中
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    doc_word = Document()
    
    # 设置页面边距（减小边距以容纳更多内容）
    sections = doc_word.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    pdf_doc = fitz.open(pdf_path)
    
    total_pages = len(pdf_doc)
    end_page = end_page if end_page else total_pages
    
    for page_num in range(start_page, min(end_page, total_pages)):
        page = pdf_doc[page_num]
        text = page.get_text("text")
        
        # 添加页码标题（紧凑样式）
        heading = doc_word.add_heading(f'━━━ 第 {page_num + 1} 页 ━━━', level=2)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # 设置标题字体大小和间距
        heading_format = heading.paragraph_format
        heading_format.space_before = Pt(0)
        heading_format.space_after = Pt(6)
        for run in heading.runs:
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(46, 125, 50)  # 绿色
        
        # 添加文本内容
        if text.strip():
            para = doc_word.add_paragraph(text)
            para_format = para.paragraph_format
            para_format.line_spacing = 1.15  # 紧凑行距
            para_format.space_before = Pt(0)
            para_format.space_after = Pt(0)
            # 设置字体大小
            for run in para.runs:
                run.font.size = Pt(10)
        else:
            para = doc_word.add_paragraph('[此页无文本内容]')
            para_format = para.paragraph_format
            para_format.space_before = Pt(0)
            para_format.space_after = Pt(0)
            for run in para.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(150, 150, 150)
        
        # 在每页内容后添加分页符（除了最后一页）
        if page_num < min(end_page, total_pages) - 1:
            doc_word.add_page_break()
    
    pdf_doc.close()
    doc_word.save(output_path)
    logger.info(f"纯文本转换完成: {end_page - start_page}页，每页PDF内容独立成页")


@app.route('/download/<filename>', methods=['GET'])
def download_file(filename):
    """
    下载转换后的文件（支持Word和PDF）
    """
    try:
        file_path = os.path.join(app.config['CONVERTED_FOLDER'], secure_filename(filename))
        
        if not os.path.exists(file_path):
            return jsonify({'error': '文件不存在'}), 404
        
        logger.info(f"下载文件: {filename}")
        
        # 根据文件扩展名设置MIME类型
        file_ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''
        mimetype_map = {
            'pdf': 'application/pdf',
            'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'doc': 'application/msword',
            'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'xls': 'application/vnd.ms-excel',
            'jpg': 'image/jpeg',
            'jpeg': 'image/jpeg',
            'png': 'image/png',
            'gif': 'image/gif',
            'bmp': 'image/bmp',
            'webp': 'image/webp'
        }
        mimetype = mimetype_map.get(file_ext, 'application/octet-stream')
        
        return send_file(
            file_path,
            as_attachment=True,
            download_name=filename,
            mimetype=mimetype
        )
    except Exception as e:
        logger.error(f"下载文件失败: {str(e)}")
        return jsonify({'error': f'下载失败: {str(e)}'}), 500


@app.route('/pdf/compress', methods=['POST'])
def compress_pdf():
    """
    PDF压缩接口
    支持三种压缩级别：low(20%), medium(50%), high(80%)
    """
    try:
        if 'file' not in request.files:
            return jsonify({'error': '未找到文件'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': '未选择文件'}), 400
        
        if not file.filename.lower().endswith('.pdf'):
            return jsonify({'error': '只支持PDF文件'}), 400
        
        # 获取压缩级别
        compression_level = request.form.get('compression_level', 'medium').lower()
        if compression_level not in ['low', 'medium', 'high']:
            compression_level = 'medium'
        
        # 压缩级别配置
        compression_configs = {
            'low': {
                'image_quality': 90,      # 图片质量90%
                'image_dpi': 200,         # 图片DPI 200
                'garbage': 2,             # 轻度清理
                'deflate': True,          # 启用压缩
                'clean': True             # 清理交叉引用
            },
            'medium': {
                'image_quality': 75,      # 图片质量75%
                'image_dpi': 150,         # 图片DPI 150
                'garbage': 3,             # 中度清理
                'deflate': True,
                'clean': True
            },
            'high': {
                'image_quality': 60,      # 图片质量60%
                'image_dpi': 100,         # 图片DPI 100
                'garbage': 4,             # 深度清理
                'deflate': True,
                'clean': True
            }
        }
        
        config = compression_configs[compression_level]
        
        # 保存上传的文件
        filename = secure_filename(file.filename)
        temp_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{uuid.uuid4().hex[:8]}_{filename}")
        file.save(temp_path)
        
        start_time = time.time()
        original_size = os.path.getsize(temp_path)
        
        logger.info(f"开始压缩PDF: {filename}, 原始大小: {original_size / 1024:.1f}KB, 压缩级别: {compression_level}")
        
        # 打开PDF
        doc = fitz.open(temp_path)
        total_pages = len(doc)
        
        # 压缩图片：遍历所有图片并压缩
        image_count = 0
        total_images = 0
        
        for page_num in range(total_pages):
            page = doc[page_num]
            image_list = page.get_images()
            total_images += len(image_list)
            
            for img_index, img in enumerate(image_list):
                try:
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]
                    
                    # 只处理JPEG和PNG图片
                    if image_ext.lower() not in ['jpeg', 'jpg', 'png']:
                        continue
                    
                    # 使用PIL压缩图片
                    img_pil = Image.open(io.BytesIO(image_bytes))
                    img_original_size = len(image_bytes)
                    
                    # 如果是PNG，转换为JPEG以减小文件大小
                    convert_to_jpeg = False
                    if image_ext.lower() == 'png':
                        if img_pil.mode in ['RGBA', 'LA']:
                            # 创建白色背景
                            background = Image.new('RGB', img_pil.size, (255, 255, 255))
                            if img_pil.mode == 'RGBA':
                                background.paste(img_pil, mask=img_pil.split()[3])
                            else:
                                background.paste(img_pil)
                            img_pil = background
                            convert_to_jpeg = True
                        elif compression_level == 'high':
                            # 高压缩级别时，所有PNG都转JPEG
                            convert_to_jpeg = True
                    
                    # 计算目标尺寸（基于DPI）
                    scale_factor = config['image_dpi'] / 72.0
                    if scale_factor < 1.0:
                        new_width = int(img_pil.width * scale_factor)
                        new_height = int(img_pil.height * scale_factor)
                        if new_width > 0 and new_height > 0:
                            img_pil = img_pil.resize((new_width, new_height), Image.Resampling.LANCZOS)
                    
                    # 压缩图片
                    img_buffer = io.BytesIO()
                    if convert_to_jpeg or image_ext.lower() in ['jpeg', 'jpg']:
                        img_pil.save(img_buffer, format='JPEG', quality=config['image_quality'], optimize=True)
                    else:
                        # PNG优化
                        img_pil.save(img_buffer, format='PNG', optimize=True)
                    
                    compressed_bytes = img_buffer.getvalue()
                    compressed_size = len(compressed_bytes)
                    
                    # 如果压缩后更小，则替换原图片
                    if compressed_size < img_original_size * 0.95:  # 至少减少5%才替换
                        # 获取图片位置信息
                        img_rects = page.get_image_rects(xref)
                        if img_rects:
                            img_rect = img_rects[0]
                            
                            # 删除原图片
                            try:
                                page.delete_image(xref)
                            except:
                                pass
                            
                            # 插入压缩后的图片
                            page.insert_image(img_rect, stream=compressed_bytes)
                            image_count += 1
                            
                except Exception as e:
                    logger.warning(f"压缩页面 {page_num + 1} 的图片 {img_index} 失败: {str(e)}")
                    continue
        
        # 生成输出文件名
        output_filename = f"compressed_{compression_level}_{uuid.uuid4().hex[:8]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        output_path = os.path.join(app.config['CONVERTED_FOLDER'], output_filename)
        
        # 保存压缩后的PDF
        doc.save(output_path,
                 garbage=config['garbage'],
                 deflate=config['deflate'],
                 clean=config['clean'])  # 新版本PyMuPDF不再支持linear参数
        
        doc.close()
        
        # 删除临时文件
        try:
            os.remove(temp_path)
        except:
            pass
        
        # 计算压缩结果
        compressed_size = os.path.getsize(output_path)
        compression_ratio = (1 - compressed_size / original_size) * 100
        elapsed_time = time.time() - start_time
        
        logger.info(f"PDF压缩完成: {output_filename}, 原始: {original_size / 1024:.1f}KB, "
                   f"压缩后: {compressed_size / 1024:.1f}KB, 压缩率: {compression_ratio:.1f}%, "
                   f"处理图片: {image_count}, 耗时: {elapsed_time:.2f}秒")
        
        return jsonify({
            'success': True,
            'filename': output_filename,
            'url': f'/download/{output_filename}',
            'original_size': original_size,
            'compressed_size': compressed_size,
            'compression_ratio': round(compression_ratio, 1),
            'compression_level': compression_level,
            'image_count': image_count,
            'total_pages': total_pages,
            'elapsed_time': round(elapsed_time, 2)
        })
        
    except Exception as e:
        logger.error(f"PDF压缩失败: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': f'压缩失败: {str(e)}'}), 500


@app.route('/pdf/arrange/upload', methods=['POST'])
def arrange_upload_pdf():
    """
    上传PDF文件用于编排（支持多文件）
    返回文件ID和页面信息
    """
    try:
        if 'file' not in request.files:
            return jsonify({'error': '未找到文件'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': '未选择文件'}), 400
        
        if not file.filename.lower().endswith('.pdf'):
            return jsonify({'error': '只支持PDF文件'}), 400
        
        # 保存文件
        filename = secure_filename(file.filename)
        file_id = uuid.uuid4().hex[:8]
        temp_path = os.path.join(app.config['UPLOAD_FOLDER'], f"arrange_{file_id}_{filename}")
        file.save(temp_path)
        
        # 打开PDF获取页面信息
        doc = fitz.open(temp_path)
        total_pages = len(doc)
        doc.close()
        
        file_size = os.path.getsize(temp_path)
        
        logger.info(f"上传PDF用于编排: {filename}, 文件ID: {file_id}, 页数: {total_pages}")
        
        return jsonify({
            'success': True,
            'file_id': file_id,
            'filename': filename,
            'file_size': file_size,
            'total_pages': total_pages
        })
        
    except Exception as e:
        logger.error(f"上传PDF失败: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': f'上传失败: {str(e)}'}), 500


@app.route('/pdf/arrange/thumbnail', methods=['POST'])
def arrange_get_thumbnail():
    """
    获取PDF指定页面的缩略图
    参数: file_id, page_number
    """
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': '缺少必要参数'}), 400
        
        file_id = data.get('file_id')
        page_number = data.get('page_number')
        
        if not file_id or not page_number:
            return jsonify({'error': '缺少file_id或page_number'}), 400
        
        # 查找对应的PDF文件
        upload_folder = app.config['UPLOAD_FOLDER']
        pdf_files = [f for f in os.listdir(upload_folder) if f.startswith(f'arrange_{file_id}_')]
        
        if not pdf_files:
            return jsonify({'error': '未找到对应的PDF文件'}), 404
        
        pdf_path = os.path.join(upload_folder, pdf_files[0])
        doc = fitz.open(pdf_path)
        
        # 验证页码
        if page_number < 1 or page_number > len(doc):
            doc.close()
            return jsonify({'error': '页码超出范围'}), 400
        
        # 生成缩略图
        page = doc[page_number - 1]
        mat = fitz.Matrix(150 / page.rect.width, 150 / page.rect.height)
        pix = page.get_pixmap(matrix=mat)
        img_data = pix.tobytes("png")
        img_base64 = base64.b64encode(img_data).decode('utf-8')
        
        doc.close()
        
        return jsonify({
            'success': True,
            'thumbnail': f'data:image/png;base64,{img_base64}'
        })
        
    except Exception as e:
        logger.error(f"获取缩略图失败: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': f'获取缩略图失败: {str(e)}'}), 500


@app.route('/file/get-pages', methods=['POST'])
def get_file_pages():
    """
    获取文件的页数（支持PDF/Word/Excel/PPT）
    """
    try:
        if 'file' not in request.files:
            return jsonify({'error': '未找到文件'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': '未选择文件'}), 400
        
        filename = secure_filename(file.filename)
        file_ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''
        
        if file_ext not in ['pdf', 'doc', 'docx', 'xls', 'xlsx', 'ppt', 'pptx']:
            return jsonify({'error': '不支持的文件格式'}), 400
        
        # 保存临时文件
        file_id = uuid.uuid4().hex[:8]
        temp_path = os.path.join(app.config['UPLOAD_FOLDER'], f"pages_{file_id}_{filename}")
        file.save(temp_path)
        
        total_pages = 0
        
        if file_ext == 'pdf':
            # PDF文件
            doc = fitz.open(temp_path)
            total_pages = len(doc)
            doc.close()
        else:
            # Office文件，需要先转PDF才能获取页数
            error_msg = None
            try:
                if not HAS_REQUESTS:
                    error_msg = '需要安装requests库才能获取Office文件页数: pip install requests'
                else:
                    java_url = 'http://localhost:8788'
                    
                    if file_ext in ['doc', 'docx']:
                        convert_url = f'{java_url}/word/topdf'
                    elif file_ext in ['xls', 'xlsx']:
                        convert_url = f'{java_url}/excel/topdf'
                    elif file_ext in ['ppt', 'pptx']:
                        convert_url = f'{java_url}/ppt/topdf'
                    else:
                        convert_url = None
                    
                    if convert_url:
                        # 检查Java服务是否可用
                        try:
                            health_check = requests.get(f'{java_url}/health', timeout=5)
                            if health_check.status_code != 200:
                                error_msg = 'Java服务(8788端口)不可用，请确保服务正在运行'
                        except requests.exceptions.RequestException:
                            error_msg = '无法连接到Java服务(8788端口)，请确保服务正在运行'
                        
                        if not error_msg:
                            with open(temp_path, 'rb') as f:
                                files = {'file': (filename, f, f'application/{file_ext}')}
                                response = requests.post(convert_url, files=files, timeout=120)
                            
                            if response.status_code == 200:
                                result = response.json()
                                if 'url' in result:
                                    # 下载PDF获取页数
                                    pdf_download_url = f"{java_url}{result['url']}"
                                    pdf_response = requests.get(pdf_download_url, timeout=60)
                                    if pdf_response.status_code == 200:
                                        pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], f"pages_{file_id}_temp.pdf")
                                        with open(pdf_path, 'wb') as pf:
                                            pf.write(pdf_response.content)
                                        doc = fitz.open(pdf_path)
                                        total_pages = len(doc)
                                        doc.close()
                                        try:
                                            os.remove(pdf_path)
                                        except:
                                            pass
                                    else:
                                        error_msg = f'下载转换后的PDF失败: {pdf_response.status_code}'
                                else:
                                    error_msg = result.get('error', '转换响应格式错误')
                            else:
                                try:
                                    error_data = response.json()
                                    error_msg = error_data.get('error', f'Java服务转换失败: {response.status_code}')
                                except:
                                    error_msg = f'Java服务转换失败: {response.status_code}'
                    else:
                        error_msg = '不支持的文件格式转换'
            except requests.exceptions.Timeout:
                error_msg = '请求超时，请检查Java服务(8788端口)是否正常运行'
            except requests.exceptions.ConnectionError:
                error_msg = '无法连接到Java服务(8788端口)，请确保服务正在运行'
            except Exception as e:
                logger.error(f"获取Office文件页数失败: {str(e)}")
                logger.error(traceback.format_exc())
                error_msg = f'获取页数失败: {str(e)}'
            
            if error_msg:
                # 删除临时文件
                try:
                    os.remove(temp_path)
                except:
                    pass
                return jsonify({
                    'success': False,
                    'error': error_msg,
                    'total_pages': 0,
                    'file_type': file_ext
                }), 500
        
        # 删除临时文件
        try:
            os.remove(temp_path)
        except:
            pass
        
        return jsonify({
            'success': True,
            'total_pages': total_pages,
            'file_type': file_ext
        })
        
    except Exception as e:
        logger.error(f"获取文件页数失败: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': f'获取页数失败: {str(e)}'}), 500


@app.route('/file/to-long-image', methods=['POST'])
def convert_to_long_image():
    """
    文件转长图（支持PDF/Word/Excel/PPT）
    1. PDF文件：直接用PyMuPDF转换
    2. Word/Excel/PPT：先转PDF，再转长图
    """
    try:
        if 'file' not in request.files:
            return jsonify({'error': '未找到文件'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': '未选择文件'}), 400
        
        filename = secure_filename(file.filename)
        file_ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''
        
        # 检查文件类型
        if file_ext not in ['pdf', 'doc', 'docx', 'xls', 'xlsx', 'ppt', 'pptx']:
            return jsonify({'error': '不支持的文件格式，仅支持PDF/Word/Excel/PPT'}), 400
        
        # 获取参数
        pages_param = request.form.get('pages', 'all')  # 页码范围
        output_format = request.form.get('format', 'jpg').lower()  # pdf/jpg/png
        dpi = int(request.form.get('dpi', 120))  # 图片DPI，降低默认值提升速度
        
        if output_format not in ['pdf', 'jpg', 'png']:
            output_format = 'jpg'
        
        # 优化DPI范围，降低上限提升速度
        if dpi < 72 or dpi > 200:
            dpi = 120
        
        start_time = time.time()
        
        # 保存上传的文件
        file_id = uuid.uuid4().hex[:8]
        temp_path = os.path.join(app.config['UPLOAD_FOLDER'], f"longimg_{file_id}_{filename}")
        file.save(temp_path)
        
        # 如果是Word/Excel/PPT，先转PDF
        pdf_path = temp_path
        if file_ext in ['doc', 'docx', 'xls', 'xlsx', 'ppt', 'pptx']:
            logger.info(f"Office文件转PDF: {filename}")
            # 调用Java服务转PDF（如果可用）
            # 或者使用Python库转PDF
            # 这里先尝试调用Java服务
            try:
                import requests
                java_url = 'http://localhost:8788'
                
                # 根据文件类型选择接口
                if file_ext in ['doc', 'docx']:
                    convert_url = f'{java_url}/word/topdf'
                elif file_ext in ['xls', 'xlsx']:
                    convert_url = f'{java_url}/excel/topdf'
                elif file_ext in ['ppt', 'pptx']:
                    convert_url = f'{java_url}/ppt/topdf'
                else:
                    convert_url = None
                
                if convert_url and HAS_REQUESTS:
                    with open(temp_path, 'rb') as f:
                        files = {'file': (filename, f, f'application/{file_ext}')}
                        response = requests.post(convert_url, files=files, timeout=120)
                    
                    if response.status_code == 200:
                        result = response.json()
                        if 'url' in result:
                            # 下载PDF文件
                            pdf_download_url = f"{java_url}{result['url']}"
                            pdf_response = requests.get(pdf_download_url, timeout=60)
                            if pdf_response.status_code == 200:
                                pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], f"longimg_{file_id}_converted.pdf")
                                with open(pdf_path, 'wb') as pf:
                                    pf.write(pdf_response.content)
                                logger.info(f"Office文件转PDF成功: {filename}")
                            else:
                                raise Exception("下载转换后的PDF失败")
                        else:
                            raise Exception("转换响应格式错误")
                    else:
                        raise Exception(f"Java服务转换失败: {response.status_code}")
                else:
                    return jsonify({'error': '不支持的文件格式转换'}), 400
                    
            except Exception as e:
                logger.warning(f"调用Java服务失败: {str(e)}")
                if not HAS_REQUESTS:
                    return jsonify({'error': '需要安装requests库才能转换Office文件: pip install requests'}), 500
                return jsonify({'error': f'Office文件转PDF失败，请确保Java服务(8788端口)正在运行: {str(e)}'}), 500
        
        # 打开PDF
        doc = fitz.open(pdf_path)
        total_pages = len(doc)
        
        # 解析页码范围
        pages_to_convert = []
        if pages_param.lower() == 'all' or pages_param.strip() == '':
            pages_to_convert = list(range(total_pages))
        else:
            # 解析页码：支持 "1,3,5" 或 "1-5" 或 "1-3,5,7-9"
            parts = re.split(r'[,，]', pages_param)
            for part in parts:
                part = part.strip()
                if '-' in part:
                    # 范围
                    start, end = part.split('-', 1)
                    start = int(start.strip())
                    end = int(end.strip())
                    for p in range(start - 1, min(end, total_pages)):
                        if p >= 0 and p not in pages_to_convert:
                            pages_to_convert.append(p)
                else:
                    # 单页
                    p = int(part.strip()) - 1
                    if 0 <= p < total_pages and p not in pages_to_convert:
                        pages_to_convert.append(p)
        
        pages_to_convert.sort()
        
        # 检查页面数量限制
        if len(pages_to_convert) > 50:
            doc.close()
            return jsonify({'error': '最多支持50页，建议20页以内'}), 400
        
        if len(pages_to_convert) == 0:
            doc.close()
            return jsonify({'error': '未选择有效页面'}), 400
        
        logger.info(f"开始转换长图: {filename}, 页数: {len(pages_to_convert)}, 格式: {output_format}, DPI: {dpi}")
        
        # 转换页面为图片（优化版）
        images = []
        gap = 0  # 页面间距（像素）
        
        # 预先计算缩放矩阵
        zoom = dpi / 72.0
        mat = fitz.Matrix(zoom, zoom)
        
        for page_idx in pages_to_convert:
            page = doc[page_idx]
            # 使用优化的pixmap参数
            pix = page.get_pixmap(matrix=mat, alpha=False)  # alpha=False提升速度
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            images.append(img)
            # 释放pixmap内存
            pix = None
        
        doc.close()
        
        # 计算长图尺寸
        max_width = max(img.width for img in images)
        total_height = sum(img.height for img in images) + gap * (len(images) - 1)
        
        # 创建长图
        long_img = Image.new('RGB', (max_width, total_height), (255, 255, 255))
        y_offset = 0
        for img in images:
            # 居中放置
            x_offset = (max_width - img.width) // 2
            long_img.paste(img, (x_offset, y_offset))
            y_offset += img.height + gap
        
        # 生成输出文件名
        output_filename = f"longimage_{file_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{output_format}"
        output_path = os.path.join(app.config['CONVERTED_FOLDER'], output_filename)
        
        # 保存长图（优化版）
        if output_format == 'pdf':
            # 将图片保存为PDF（使用PyMuPDF）
            pdf_doc = fitz.open()
            # 创建页面（尺寸转换为点，PyMuPDF使用点为单位）
            page = pdf_doc.new_page(width=max_width, height=total_height)
            # 将PIL图片转换为bytes（使用JPEG格式更快）
            img_bytes = io.BytesIO()
            long_img.save(img_bytes, format='JPEG', quality=85)
            img_bytes.seek(0)
            # 插入图片到PDF
            page.insert_image(fitz.Rect(0, 0, max_width, total_height), stream=img_bytes.getvalue())
            # 移除clean参数提升保存速度
            pdf_doc.save(output_path, garbage=4, deflate=True)
            pdf_doc.close()
        elif output_format == 'jpg':
            # 降低质量提升速度，85质量已足够
            long_img.save(output_path, format='JPEG', quality=85, optimize=True)
        else:  # png
            # PNG使用压缩级别控制
            long_img.save(output_path, format='PNG', compress_level=6, optimize=False)
        
        # 删除临时文件
        try:
            if pdf_path != temp_path:
                os.remove(pdf_path)
            os.remove(temp_path)
        except:
            pass
        
        output_size = os.path.getsize(output_path)
        elapsed_time = time.time() - start_time
        
        logger.info(f"长图转换完成: {output_filename}, 页数: {len(pages_to_convert)}, 尺寸: {max_width}x{total_height}, "
                   f"格式: {output_format}, 大小: {output_size / 1024:.1f}KB, 耗时: {elapsed_time:.2f}秒")
        
        return jsonify({
            'success': True,
            'filename': output_filename,
            'url': f'/download/{output_filename}',
            'format': output_format,
            'width': max_width,
            'height': total_height,
            'file_size': output_size,
            'pages_count': len(pages_to_convert),
            'elapsed_time': round(elapsed_time, 2)
        })
        
    except Exception as e:
        logger.error(f"转长图失败: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': f'转换失败: {str(e)}'}), 500


# 辅助：加载中文字体（用于拼图描述）
def load_font(size=36):
    """尝试加载系统中文字体，失败则回退默认字体"""
    font_candidates = [
        r"C:\\Windows\\Fonts\\msyh.ttc",
        r"C:\\Windows\\Fonts\\simhei.ttf",
        r"C:\\Windows\\Fonts\\simsun.ttc",
        "msyh.ttc",
        "simhei.ttf",
        "simsun.ttc"
    ]
    for path in font_candidates:
        try:
            if os.path.exists(path):
                return ImageFont.truetype(path, size=size)
        except Exception:
            continue
    return ImageFont.load_default()


@app.route('/image/collage', methods=['POST'])
def image_collage():
    """
    图片拼图
    - 支持A4画布（300dpi，2480x3508）
    - 支持6种布局：
        full           : 单图铺满
        two_vertical   : 左右两列
        three_vertical : 三列
        two_horizontal : 上下两行
        four_grid      : 2x2
        six_grid       : 2x3
    - captions: 每张图片下方15字描述
    - output_type: image（jpg）或 pdf（多页）
    """
    try:
        start_time = time.time()
        data = request.get_json()
        if not data:
            return jsonify({'error': '缺少必要参数'}), 400

        layout = data.get('layout', 'full')
        output_type = data.get('output_type', 'image')
        images_base64 = data.get('images', [])
        captions = data.get('captions', [])

        if not images_base64:
            return jsonify({'error': '请至少上传一张图片'}), 400

        # 限制数量，避免内存压力
        if len(images_base64) > 30:
            return jsonify({'error': '最多支持30张图片'}), 400

        # A4尺寸（300dpi）
        A4_WIDTH, A4_HEIGHT = 2480, 3508
        background_color = (255, 255, 255)
        caption_height = 100  # 预留描述高度
        caption_font = load_font(size=36)
        caption_color = (80, 80, 80)

        # 布局定义：归一化坐标（0-1）
        layouts = {
            'full': [
                {'x': 0.04, 'y': 0.04, 'w': 0.92, 'h': 0.92}
            ],
            'two_vertical': [
                {'x': 0.03, 'y': 0.04, 'w': 0.47, 'h': 0.92},
                {'x': 0.50, 'y': 0.04, 'w': 0.47, 'h': 0.92}
            ],
            'three_vertical': [
                {'x': 0.02, 'y': 0.04, 'w': 0.31, 'h': 0.92},
                {'x': 0.345, 'y': 0.04, 'w': 0.31, 'h': 0.92},
                {'x': 0.67, 'y': 0.04, 'w': 0.31, 'h': 0.92}
            ],
            'two_horizontal': [
                {'x': 0.04, 'y': 0.04, 'w': 0.92, 'h': 0.45},
                {'x': 0.04, 'y': 0.51, 'w': 0.92, 'h': 0.45}
            ],
            'four_grid': [
                {'x': 0.04, 'y': 0.04, 'w': 0.44, 'h': 0.44},
                {'x': 0.52, 'y': 0.04, 'w': 0.44, 'h': 0.44},
                {'x': 0.04, 'y': 0.52, 'w': 0.44, 'h': 0.44},
                {'x': 0.52, 'y': 0.52, 'w': 0.44, 'h': 0.44}
            ],
            'six_grid': [
                {'x': 0.04, 'y': 0.04, 'w': 0.44, 'h': 0.28},
                {'x': 0.52, 'y': 0.04, 'w': 0.44, 'h': 0.28},
                {'x': 0.04, 'y': 0.36, 'w': 0.44, 'h': 0.28},
                {'x': 0.52, 'y': 0.36, 'w': 0.44, 'h': 0.28},
                {'x': 0.04, 'y': 0.68, 'w': 0.44, 'h': 0.28},
                {'x': 0.52, 'y': 0.68, 'w': 0.44, 'h': 0.28}
            ]
        }

        if layout not in layouts:
            layout = 'full'

        slot_defs = layouts[layout]
        slots_per_page = len(slot_defs)
        total_images = len(images_base64)

        if output_type == 'image' and total_images > slots_per_page:
            logger.info("拼图导出图片模式：超过版位的图片将被忽略")
            images_base64 = images_base64[:slots_per_page]
            captions = captions[:slots_per_page]
            total_images = len(images_base64)

        page_images = []
        warnings = []

        # 分页处理（用于PDF多页）
        chunk_start = 0
        page_index = 0
        while chunk_start < total_images:
            chunk_end = min(chunk_start + slots_per_page, total_images)
            chunk = images_base64[chunk_start:chunk_end]
            chunk_captions = captions[chunk_start:chunk_end] if captions else []

            canvas = Image.new('RGB', (A4_WIDTH, A4_HEIGHT), background_color)
            draw = ImageDraw.Draw(canvas)

            for idx, img_b64 in enumerate(chunk):
                try:
                    # 处理dataURL
                    if ',' in img_b64:
                        img_b64 = img_b64.split(',', 1)[1]
                    img_bytes = base64.b64decode(img_b64)
                    img = Image.open(io.BytesIO(img_bytes)).convert('RGB')

                    slot = slot_defs[idx]
                    x = int(slot['x'] * A4_WIDTH)
                    y = int(slot['y'] * A4_HEIGHT)
                    w = int(slot['w'] * A4_WIDTH)
                    h = int(slot['h'] * A4_HEIGHT)

                    # 为描述预留空间
                    img_h = h - caption_height
                    if img_h <= 0:
                        img_h = h

                    # 等比缩放并居中
                    scale = min(w / img.width, img_h / img.height)
                    new_w = max(1, int(img.width * scale))
                    new_h = max(1, int(img.height * scale))
                    img_resized = img.resize((new_w, new_h), resample=Image.LANCZOS)
                    paste_x = x + (w - new_w) // 2
                    paste_y = y + (img_h - new_h) // 2
                    canvas.paste(img_resized, (paste_x, paste_y))

                    # 描述文字
                    caption = ''
                    if idx < len(chunk_captions):
                        caption = (chunk_captions[idx] or '')[:15]
                    if caption:
                        text_w, text_h = draw.textsize(caption, font=caption_font)
                        text_x = x + (w - text_w) // 2
                        text_y = y + img_h + (caption_height - text_h) // 2
                        draw.text((text_x, text_y), caption, font=caption_font, fill=caption_color)

                except Exception as e:
                    logger.warning(f"处理第{idx+1}张图片失败: {e}")
                    warnings.append(f"第{idx+1}张图片处理失败: {e}")
                    continue

            page_images.append(canvas)
            chunk_start = chunk_end
            page_index += 1

        if not page_images:
            return jsonify({'error': '图片处理失败，请检查上传的图片格式'}), 400

        timestamp = int(time.time() * 1000)
        if output_type == 'pdf':
            output_filename = f"collage_{timestamp}.pdf"
            output_path = os.path.join(app.config['CONVERTED_FOLDER'], output_filename)

            pdf_doc = fitz.open()
            for page_img in page_images:
                img_bytes = io.BytesIO()
                page_img.save(img_bytes, format='JPEG', quality=90, optimize=True)
                img_bytes.seek(0)
                page = pdf_doc.new_page(width=A4_WIDTH, height=A4_HEIGHT)
                page.insert_image(fitz.Rect(0, 0, A4_WIDTH, A4_HEIGHT), stream=img_bytes.getvalue())
            pdf_doc.save(output_path, garbage=4, deflate=True)
            pdf_doc.close()
            result_format = 'pdf'
        else:
            # 单页图片输出（使用第一页）
            output_filename = f"collage_{timestamp}.jpg"
            output_path = os.path.join(app.config['CONVERTED_FOLDER'], output_filename)
            page_images[0].save(output_path, format='JPEG', quality=90, optimize=True)
            result_format = 'jpg'

        output_size = os.path.getsize(output_path)
        elapsed_time = time.time() - start_time

        logger.info(f"拼图完成: {output_filename}, 页数: {len(page_images)}, 耗时: {elapsed_time:.2f}s, 大小: {output_size/1024:.1f}KB")

        return jsonify({
            'success': True,
            'filename': output_filename,
            'url': f'/download/{output_filename}',
            'format': result_format,
            'pages': len(page_images),
            'file_size': output_size,
            'elapsed_time': round(elapsed_time, 2),
            'warning': warnings if warnings else None
        })

    except Exception as e:
        logger.error(f"拼图失败: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': f'拼图失败: {str(e)}'}), 500


@app.route('/pdf/arrange/generate', methods=['POST'])
def arrange_generate_pdf():
    """
    根据用户选择的页面顺序生成编排后的PDF
    参数: files (JSON数组，每个元素包含file_id和pages数组)
    """
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': '缺少必要参数'}), 400
        
        # 接收页面顺序数组：[{file_id: 'xxx', page_number: 1}, ...]
        pages_order = data.get('pages_order', [])
        
        # 兼容旧格式：files数组
        if not pages_order and 'files' in data:
            files_config = data['files']
            # 转换为pages_order格式
            pages_order = []
            for file_config in files_config:
                file_id = file_config.get('file_id')
                pages = file_config.get('pages', [])
                for page_num in pages:
                    pages_order.append({
                        'file_id': file_id,
                        'page_number': page_num
                    })
        
        if not pages_order or len(pages_order) == 0:
            return jsonify({'error': '未选择任何页面'}), 400
        
        start_time = time.time()
        
        # 创建新的PDF文档
        output_doc = fitz.open()
        
        total_pages_added = 0
        
        # 按用户选择的顺序添加页面
        file_cache = {}  # 缓存已打开的文档
        
        for page_info in pages_order:
            file_id = page_info.get('file_id')
            page_number = page_info.get('page_number')
            
            if not file_id or not page_number:
                continue
            
            # 从缓存获取或打开PDF文件
            if file_id not in file_cache:
                upload_folder = app.config['UPLOAD_FOLDER']
                pdf_files = [f for f in os.listdir(upload_folder) if f.startswith(f'arrange_{file_id}_')]
                
                if not pdf_files:
                    logger.warning(f"未找到文件ID对应的PDF: {file_id}")
                    continue
                
                pdf_path = os.path.join(upload_folder, pdf_files[0])
                file_cache[file_id] = fitz.open(pdf_path)
            
            source_doc = file_cache[file_id]
            
            # 添加指定页面
            if 1 <= page_number <= len(source_doc):
                page_idx = page_number - 1
                # 复制页面到新文档
                output_doc.insert_pdf(source_doc, from_page=page_idx, to_page=page_idx)
                total_pages_added += 1
        
        # 关闭所有打开的文档
        for doc in file_cache.values():
            doc.close()
        
        if total_pages_added == 0:
            output_doc.close()
            return jsonify({'error': '未添加任何页面'}), 400
        
        # 生成输出文件名
        output_filename = f"arranged_{uuid.uuid4().hex[:8]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        output_path = os.path.join(app.config['CONVERTED_FOLDER'], output_filename)
        
        # 保存PDF
        output_doc.save(output_path, garbage=4, deflate=True, clean=True)
        output_doc.close()
        
        output_size = os.path.getsize(output_path)
        elapsed_time = time.time() - start_time
        
        logger.info(f"PDF编排完成: {output_filename}, 总页数: {total_pages_added}, 耗时: {elapsed_time:.2f}秒")
        
        return jsonify({
            'success': True,
            'filename': output_filename,
            'url': f'/download/{output_filename}',
            'total_pages': total_pages_added,
            'file_size': output_size,
            'elapsed_time': round(elapsed_time, 2)
        })
        
    except Exception as e:
        logger.error(f"PDF编排失败: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': f'编排失败: {str(e)}'}), 500


@app.route('/longimage/print', methods=['POST'])
def print_longimage():
    """
    长图打印功能：将长图裁剪并排版到A4页面中
    支持1/2/3列布局
    输入：图片或PDF文件
    输出：适合打印的PDF文件
    """
    try:
        # 检查文件
        if 'file' not in request.files:
            return jsonify({'error': '未上传文件'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': '文件名为空'}), 400
        
        # 获取分栏选项（1/2/3列）
        columns = int(request.form.get('columns', 1))
        if columns not in [1, 2, 3]:
            return jsonify({'error': '分栏数必须是1、2或3'}), 400
        
        # 获取DPI设置（影响质量）
        dpi = int(request.form.get('dpi', 300))
        
        start_time = time.time()
        
        # 保存上传的文件
        filename = secure_filename(file.filename)
        file_ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
        upload_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{uuid.uuid4()}_{filename}")
        file.save(upload_path)
        
        logger.info(f"开始处理长图打印: {filename}, 分栏: {columns}, DPI: {dpi}")
        
        # 将输入转换为PIL Image
        images = []
        if file_ext == 'pdf':
            # PDF文件：转换每一页为图片
            pdf_doc = fitz.open(upload_path)
            for page_num in range(len(pdf_doc)):
                page = pdf_doc[page_num]
                # 使用指定DPI渲染
                mat = fitz.Matrix(dpi/72, dpi/72)
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.tobytes("png")
                img = Image.open(io.BytesIO(img_data))
                images.append(img)
            pdf_doc.close()
        else:
            # 图片文件：直接打开
            img = Image.open(upload_path)
            # 转换为RGB（处理RGBA等格式）
            if img.mode != 'RGB':
                img = img.convert('RGB')
            images.append(img)
        
        # A4尺寸（72 DPI下的点数）
        A4_WIDTH_PT = 595
        A4_HEIGHT_PT = 842
        
        # 页面边距（点）
        MARGIN = 20
        
        # 可用区域尺寸
        usable_width = A4_WIDTH_PT - 2 * MARGIN
        usable_height = A4_HEIGHT_PT - 2 * MARGIN
        
        # 计算每列的宽度
        column_spacing = 10  # 列间距
        if columns == 1:
            column_width = usable_width
        elif columns == 2:
            column_width = (usable_width - column_spacing) / 2
        else:  # 3列
            column_width = (usable_width - 2 * column_spacing) / 3
        
        # 创建输出PDF
        output_pdf = fitz.open()
        
        # 收集所有需要处理的图片段
        all_segments = []
        
        for img_index, source_img in enumerate(images):
            logger.info(f"处理第 {img_index + 1}/{len(images)} 张图片, 尺寸: {source_img.size}")
            
            # 计算缩放比例（保持宽高比）
            img_width, img_height = source_img.size
            scale = (column_width * dpi / 72) / img_width
            
            # 调整图片尺寸以适应列宽
            new_width = int(img_width * scale)
            new_height = int(img_height * scale)
            resized_img = source_img.resize((new_width, new_height), Image.Resampling.LANCZOS)
            
            logger.info(f"图片缩放后尺寸: {new_width}x{new_height}px")
            
            # 将整张图片添加到待处理列表（不再预先切割）
            # 将整个图片添加到待处理列表
            all_segments.append({
                'image': resized_img,
                'width': new_width,
                'height': new_height
            })
        
        # 智能多列布局算法
        logger.info(f"开始智能布局: 共{len(all_segments)}张图片, {columns}列")
        
        # 初始化页面状态
        current_page = None
        current_column = 0  # 当前列索引（0, 1, 2...）
        column_y_positions = [MARGIN] * columns  # 每列的当前Y位置
        
        # 遍历所有图片段
        for seg_data in all_segments:
            img = seg_data['image']
            img_width_px = seg_data['width']
            img_height_px = seg_data['height']
            
            # 需要切割成多段以适应列高
            segment_height_px = int(usable_height * dpi / 72)
            remaining_height = img_height_px
            y_offset = 0
            
            while remaining_height > 0:
                # 计算当前段的高度
                current_seg_height = min(segment_height_px, remaining_height)
                
                # 裁剪当前段
                segment = img.crop((0, y_offset, img_width_px, y_offset + current_seg_height))
                
                # 计算segment在PDF中的尺寸（点）
                seg_width_pt = column_width
                seg_height_pt = current_seg_height * 72 / dpi
                
                # 查找可以放置的列
                placed = False
                for attempt in range(columns):
                    check_col = (current_column + attempt) % columns
                    
                    # 检查当前列是否有足够空间
                    if column_y_positions[check_col] + seg_height_pt <= A4_HEIGHT_PT - MARGIN:
                        # 如果需要新页面
                        if current_page is None:
                            current_page = output_pdf.new_page(width=A4_WIDTH_PT, height=A4_HEIGHT_PT)
                            column_y_positions = [MARGIN] * columns
                        
                        # 计算X位置
                        if columns == 1:
                            x = MARGIN
                        elif columns == 2:
                            x = MARGIN + check_col * (column_width + column_spacing)
                        else:  # 3列
                            x = MARGIN + check_col * (column_width + column_spacing)
                        
                        y = column_y_positions[check_col]
                        
                        # 保存segment为临时文件
                        temp_seg_path = os.path.join(app.config['UPLOAD_FOLDER'], f"temp_seg_{uuid.uuid4()}.png")
                        segment.save(temp_seg_path, "PNG", optimize=True, quality=95)
                        
                        # 插入图片到PDF
                        rect = fitz.Rect(x, y, x + seg_width_pt, y + seg_height_pt)
                        current_page.insert_image(rect, filename=temp_seg_path)
                        
                        # 删除临时文件
                        try:
                            os.remove(temp_seg_path)
                        except:
                            pass
                        
                        # 更新该列的Y位置
                        column_y_positions[check_col] += seg_height_pt
                        current_column = check_col
                        placed = True
                        break
                
                # 如果所有列都满了，创建新页面
                if not placed:
                    current_page = output_pdf.new_page(width=A4_WIDTH_PT, height=A4_HEIGHT_PT)
                    column_y_positions = [MARGIN] * columns
                    current_column = 0
                    
                    # 在新页面的第一列放置
                    x = MARGIN
                    y = MARGIN
                    
                    temp_seg_path = os.path.join(app.config['UPLOAD_FOLDER'], f"temp_seg_{uuid.uuid4()}.png")
                    segment.save(temp_seg_path, "PNG", optimize=True, quality=95)
                    
                    rect = fitz.Rect(x, y, x + seg_width_pt, y + seg_height_pt)
                    current_page.insert_image(rect, filename=temp_seg_path)
                    
                    try:
                        os.remove(temp_seg_path)
                    except:
                        pass
                    
                    column_y_positions[0] = MARGIN + seg_height_pt
                
                # 更新剩余高度和偏移
                remaining_height -= current_seg_height
                y_offset += current_seg_height
        
        # 获取页数（在关闭之前）
        total_pages = len(output_pdf)
        
        # 保存输出PDF
        output_filename = f"print_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}.pdf"
        output_path = os.path.join(app.config['CONVERTED_FOLDER'], output_filename)
        output_pdf.save(output_path, garbage=4, deflate=True)
        output_pdf.close()
        
        # 清理上传的文件
        try:
            os.remove(upload_path)
        except:
            pass
        
        # 获取输出文件大小
        output_size = os.path.getsize(output_path)
        elapsed_time = time.time() - start_time
        
        logger.info(f"长图打印完成: {output_filename}, 页数: {total_pages}, 耗时: {elapsed_time:.2f}秒")
        
        return jsonify({
            'success': True,
            'filename': output_filename,
            'url': f'/download/{output_filename}',
            'total_pages': total_pages,
            'file_size': output_size,
            'elapsed_time': round(elapsed_time, 2),
            'columns': columns
        })
        
    except Exception as e:
        logger.error(f"长图打印失败: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': f'处理失败: {str(e)}'}), 500


@app.route('/file/to-grayscale', methods=['POST'])
def convert_to_grayscale():
    """
    将PDF或图片转换为黑白/灰度图
    支持两种模式：
    1. gray - 仅去除彩色（灰度图）
    2. bw - 转换为白底黑字（二值化）
    """
    try:
        # 检查文件
        if 'file' not in request.files:
            return jsonify({'error': '未上传文件'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': '文件名为空'}), 400
        
        # 获取参数
        mode = request.form.get('mode', 'gray')  # 'gray' 或 'bw'
        if mode not in ['gray', 'bw']:
            mode = 'gray'
        
        start_time = time.time()
        
        # 保存上传的文件
        filename = secure_filename(file.filename)
        file_ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
        upload_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{uuid.uuid4()}_{filename}")
        file.save(upload_path)
        
        logger.info(f"开始转换黑白: {filename}, 模式: {mode}")
        
        # 判断文件类型
        is_pdf = file_ext == 'pdf'
        
        if is_pdf:
            # 处理PDF文件
            doc = fitz.open(upload_path)
            total_pages = len(doc)
            
            # 限制页数
            if total_pages > 20:
                doc.close()
                os.remove(upload_path)
                return jsonify({'error': f'PDF页数超过20页（当前{total_pages}页），请使用文件拆分功能'}), 400
            
            # 创建输出PDF
            output_pdf = fitz.open()
            
            for page_num in range(total_pages):
                page = doc[page_num]
                
                # 使用超高分辨率渲染（确保文字清晰）
                # 白底黑字模式使用更高的DPI以获得最佳效果
                if mode == 'bw':
                    zoom_factor = 4  # 288 DPI，高质量
                else:
                    zoom_factor = 3  # 216 DPI，灰度模式
                
                mat = fitz.Matrix(zoom_factor, zoom_factor)
                pix = page.get_pixmap(matrix=mat, colorspace=fitz.csGRAY, alpha=False)
                
                # 获取原始页面尺寸
                orig_width = page.rect.width
                orig_height = page.rect.height
                
                # 如果是黑白模式，进行二值化处理
                if mode == 'bw':
                    # 使用PIL进行二值化
                    img = Image.frombytes("L", [pix.width, pix.height], pix.samples)
                    
                    # OTSU自动阈值
                    import numpy as np
                    img_array = np.array(img)
                    hist, _ = np.histogram(img_array.flatten(), 256, [0, 256])
                    total = img_array.size
                    sum_total = np.sum(np.arange(256) * hist)
                    sum_bg = 0
                    weight_bg = 0
                    max_variance = 0
                    threshold = 0
                    
                    for i in range(256):
                        weight_bg += hist[i]
                        if weight_bg == 0:
                            continue
                        weight_fg = total - weight_bg
                        if weight_fg == 0:
                            break
                        sum_bg += i * hist[i]
                        mean_bg = sum_bg / weight_bg
                        mean_fg = (sum_total - sum_bg) / weight_fg
                        variance = weight_bg * weight_fg * (mean_bg - mean_fg) ** 2
                        if variance > max_variance:
                            max_variance = variance
                            threshold = i
                    
                    # 稍微降低阈值让字体更粗
                    adjusted_threshold = max(0, threshold - 10)
                    img = img.point(lambda x: 255 if x > adjusted_threshold else 0, '1')
                    img = img.convert('L')
                    
                    # 保存为临时PNG文件（高质量）
                    temp_img_path = os.path.join(app.config['UPLOAD_FOLDER'], f"temp_bw_{uuid.uuid4()}.png")
                    img.save(temp_img_path, format='PNG', compress_level=0, dpi=(zoom_factor * 72, zoom_factor * 72))
                    
                    # 创建新页面（保持原始尺寸）
                    new_page = output_pdf.new_page(width=orig_width, height=orig_height)
                    
                    # 插入图片，指定DPI信息
                    new_page.insert_image(
                        fitz.Rect(0, 0, orig_width, orig_height),
                        filename=temp_img_path
                    )
                    
                    # 删除临时文件
                    try:
                        os.remove(temp_img_path)
                    except:
                        pass
                else:
                    # 灰度模式：直接使用pixmap
                    new_page = output_pdf.new_page(width=orig_width, height=orig_height)
                    
                    # 保存pixmap为临时PNG
                    temp_img_path = os.path.join(app.config['UPLOAD_FOLDER'], f"temp_gray_{uuid.uuid4()}.png")
                    pix.save(temp_img_path)
                    
                    # 插入图片
                    new_page.insert_image(
                        fitz.Rect(0, 0, orig_width, orig_height),
                        filename=temp_img_path
                    )
                    
                    # 删除临时文件
                    try:
                        os.remove(temp_img_path)
                    except:
                        pass
                
                pix = None  # 释放内存
            
            doc.close()
            
            # 保存输出PDF（使用最高质量设置）
            output_filename = f"grayscale_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}.pdf"
            output_path = os.path.join(app.config['CONVERTED_FOLDER'], output_filename)
            # 使用deflate压缩但不降低图片质量
            # garbage=4 进行垃圾回收，deflate=True 压缩但保持质量
            output_pdf.save(
                output_path, 
                garbage=4,        # 垃圾回收
                deflate=True,     # 压缩
                no_new_id=True    # 不生成新ID，保持稳定性
            )
            output_pdf.close()
            
        else:
            # 处理图片文件
            img = Image.open(upload_path)
            
            # 转换为灰度
            gray_img = img.convert('L')
            
            # 如果是黑白模式，进行二值化
            if mode == 'bw':
                # 使用OTSU自动阈值
                import numpy as np
                img_array = np.array(gray_img)
                hist, _ = np.histogram(img_array.flatten(), 256, [0, 256])
                total = img_array.size
                sum_total = np.sum(np.arange(256) * hist)
                sum_bg = 0
                weight_bg = 0
                max_variance = 0
                threshold = 0
                
                for i in range(256):
                    weight_bg += hist[i]
                    if weight_bg == 0:
                        continue
                    weight_fg = total - weight_bg
                    if weight_fg == 0:
                        break
                    sum_bg += i * hist[i]
                    mean_bg = sum_bg / weight_bg
                    mean_fg = (sum_total - sum_bg) / weight_fg
                    variance = weight_bg * weight_fg * (mean_bg - mean_fg) ** 2
                    if variance > max_variance:
                        max_variance = variance
                        threshold = i
                
                # 应用阈值（稍微降低阈值让字体更粗）
                adjusted_threshold = max(0, threshold - 15)
                gray_img = gray_img.point(lambda x: 255 if x > adjusted_threshold else 0, '1')
                gray_img = gray_img.convert('L')
            
            # 保存输出图片
            output_filename = f"grayscale_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}.{file_ext}"
            output_path = os.path.join(app.config['CONVERTED_FOLDER'], output_filename)
            gray_img.save(output_path, quality=95, optimize=True)
        
        # 清理上传文件
        try:
            os.remove(upload_path)
        except:
            pass
        
        # 获取输出文件信息
        output_size = os.path.getsize(output_path)
        elapsed_time = time.time() - start_time
        
        logger.info(f"黑白转换完成: {output_filename}, 模式: {mode}, 耗时: {elapsed_time:.2f}秒")
        
        return jsonify({
            'success': True,
            'filename': output_filename,
            'url': f'/download/{output_filename}',
            'file_size': output_size,
            'mode': mode,
            'elapsed_time': round(elapsed_time, 2),
            'pages': total_pages if is_pdf else 1
        })
        
    except Exception as e:
        logger.error(f"转换黑白失败: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': f'转换失败: {str(e)}'}), 500


@app.route('/file/rename', methods=['POST'])
def rename_file():
    """
    文件重命名功能
    支持所有格式的文件
    复制文件后重命名，不影响原文件
    """
    try:
        # 检查是否有文件
        if 'file' not in request.files:
            logger.warning("重命名接口：未找到上传的文件")
            return jsonify({'error': '请选择要重命名的文件'}), 400
        
        file = request.files['file']
        if file.filename == '':
            logger.warning("重命名接口：文件名为空")
            return jsonify({'error': '请选择要重命名的文件'}), 400
        
        # 获取新文件名
        new_filename = request.form.get('new_filename', '').strip()
        if not new_filename:
            logger.warning("重命名接口：新文件名为空")
            return jsonify({'error': '请输入新的文件名'}), 400
        
        # 获取原始文件名和扩展名
        original_filename = secure_filename(file.filename)
        _, original_ext = os.path.splitext(original_filename)
        
        # 处理新文件名：如果用户没有提供扩展名，使用原文件的扩展名
        if not os.path.splitext(new_filename)[1]:
            # 新文件名没有扩展名，添加原扩展名
            new_filename_with_ext = new_filename + original_ext
        else:
            # 新文件名已有扩展名，直接使用
            new_filename_with_ext = new_filename
        
        # 确保文件名安全
        safe_new_filename = secure_filename(new_filename_with_ext)
        
        # 生成唯一的文件名（避免冲突）
        timestamp = int(time.time() * 1000)
        unique_filename = f"{timestamp}_{safe_new_filename}"
        
        # 保存上传的文件到临时位置
        temp_original_path = os.path.join(app.config['UPLOAD_FOLDER'], f"temp_{timestamp}_{original_filename}")
        file.save(temp_original_path)
        
        logger.info(f"文件重命名: {original_filename} -> {safe_new_filename}")
        
        # 复制并重命名文件
        renamed_file_path = os.path.join(app.config['CONVERTED_FOLDER'], unique_filename)
        shutil.copy2(temp_original_path, renamed_file_path)
        
        # 删除临时文件
        try:
            os.remove(temp_original_path)
        except Exception as e:
            logger.warning(f"删除临时文件失败: {e}")
        
        # 获取文件信息
        file_size = os.path.getsize(renamed_file_path)
        file_size_mb = file_size / (1024 * 1024)
        
        # 返回结果
        result = {
            'success': True,
            'message': '文件重命名成功',
            'original_name': original_filename,
            'new_name': safe_new_filename,
            'download_filename': unique_filename,
            'file_size': file_size,
            'file_size_mb': round(file_size_mb, 2)
        }
        
        logger.info(f"文件重命名成功: {unique_filename} ({file_size_mb:.2f}MB)")
        return jsonify(result)
        
    except Exception as e:
        logger.error(f"文件重命名失败: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': f'文件重命名失败: {str(e)}'}), 500


@app.route('/pdf/add-page-numbers', methods=['POST'])
def add_page_numbers():
    """
    为PDF添加页码功能
    支持设置页码位置、样式、大小、边距、颜色、起始页码
    高性能、高质量、快速处理
    """
    try:
        start_time = time.time()
        
        # 检查是否有文件
        if 'file' not in request.files:
            logger.warning("添加页码接口：未找到上传的文件")
            return jsonify({'error': '请选择要添加页码的PDF文件'}), 400
        
        file = request.files['file']
        if file.filename == '':
            logger.warning("添加页码接口：文件名为空")
            return jsonify({'error': '请选择要添加页码的PDF文件'}), 400
        
        # 检查文件大小（200MB限制）
        file.seek(0, os.SEEK_END)
        file_size = file.tell()
        file.seek(0)
        if file_size > 200 * 1024 * 1024:
            return jsonify({'error': '文件大小不能超过200MB'}), 400
        
        # 获取参数
        position = request.form.get('position', 'bottom_center')  # 位置：top_left, top_center, top_right, bottom_left, bottom_center, bottom_right
        style = request.form.get('style', 'simple')  # 样式：simple(1,2,3), fraction(1/10,2/10), roman(i,ii,iii)
        font_size = int(request.form.get('font_size', '12'))
        margin = int(request.form.get('margin', '20'))  # 边距（像素）
        color = request.form.get('color', '#000000')  # 颜色，格式：#RRGGBB
        start_page = int(request.form.get('start_page', '1'))  # 从第几页开始添加页码
        
        # 验证参数
        if font_size < 1 or font_size > 200:
            return jsonify({'error': '字体大小必须在1-200之间'}), 400
        if margin < 0 or margin > 200:
            return jsonify({'error': '边距必须在0-200之间'}), 400
        if start_page < 1:
            return jsonify({'error': '起始页码必须大于0'}), 400
        
        # 解析颜色
        try:
            color_rgb = tuple(int(color[i:i+2], 16) for i in (1, 3, 5))  # #RRGGBB -> (R, G, B)
            color_normalized = tuple(c / 255.0 for c in color_rgb)  # 归一化到0-1
        except:
            color_normalized = (0, 0, 0)  # 默认黑色
        
        # 保存上传的文件
        timestamp = int(time.time() * 1000)
        original_filename = secure_filename(file.filename)
        temp_input_path = os.path.join(app.config['UPLOAD_FOLDER'], f"temp_pagenum_{timestamp}_{original_filename}")
        file.save(temp_input_path)
        
        logger.info(f"开始添加页码: {original_filename}, 位置={position}, 样式={style}, 大小={font_size}, 起始页={start_page}")
        
        # 打开PDF
        doc = fitz.open(temp_input_path)
        total_pages = len(doc)
        
        if start_page > total_pages:
            doc.close()
            os.remove(temp_input_path)
            return jsonify({'error': f'起始页码({start_page})不能大于总页数({total_pages})'}), 400
        
        # 遍历每一页添加页码
        for page_num in range(total_pages):
            page = doc[page_num]
            rect = page.rect
            
            # 只从start_page开始添加页码
            if page_num + 1 < start_page:
                continue
            
            # 计算页码文本（从start_page开始计数，从1开始）
            # 例如：如果start_page=3，那么第3页显示页码1，第4页显示页码2，以此类推
            actual_page_num = (page_num + 1) - start_page + 1
            # 计算实际需要添加页码的总页数（用于样式中的总页数显示）
            actual_total_pages = total_pages - start_page + 1
            page_text = format_page_number(actual_page_num, actual_total_pages, style)
            
            # 计算页码位置
            if position == 'top_left':
                x = margin
                y = margin
                align = 0  # 左对齐
            elif position == 'top_center':
                x = rect.width / 2
                y = margin
                align = 1  # 居中
            elif position == 'top_right':
                x = rect.width - margin
                y = margin
                align = 2  # 右对齐
            elif position == 'bottom_left':
                x = margin
                y = rect.height - margin
                align = 0
            elif position == 'bottom_center':
                x = rect.width / 2
                y = rect.height - margin
                align = 1
            elif position == 'bottom_right':
                x = rect.width - margin
                y = rect.height - margin
                align = 2
            else:
                # 默认底部居中
                x = rect.width / 2
                y = rect.height - margin
                align = 1
            
            # 使用insert_text，根据对齐方式手动计算文本位置
            # 更精确地估算文本宽度（区分中文字符和数字/字母）
            text_width = 0
            for char in page_text:
                # 判断是否为中文字符（包括中文标点）
                if '\u4e00' <= char <= '\u9fff' or char in '，。、；：？！':
                    text_width += font_size  # 中文字符宽度约为字体大小
                else:
                    text_width += font_size * 0.55  # 数字和字母宽度约为字体大小的0.55倍
            
            # 根据对齐方式调整x坐标
            # 增加一些宽度余量，确保文本完全显示（避免省略号）
            text_width_with_margin = text_width * 1.1  # 增加10%的余量
            
            if align == 1:  # 居中对齐
                # x已经是中心点，需要减去文本宽度的一半
                text_x = x - text_width_with_margin / 2
            elif align == 2:  # 右对齐
                # x是右边界，需要减去文本宽度
                text_x = x - text_width_with_margin
            else:  # 左对齐 (align == 0)
                # x就是左边界，直接使用
                text_x = x
            
            # 确保文本不超出页面边界
            if text_x < 0:
                text_x = 0
            elif text_x + text_width_with_margin > rect.width:
                text_x = rect.width - text_width_with_margin
                if text_x < 0:
                    text_x = 0
            
            # 添加页码文本
            # 检查文本是否包含中文字符，如果包含则使用支持中文的字体
            has_chinese = any('\u4e00' <= char <= '\u9fff' for char in page_text)
            
            text_inserted = False
            
            if has_chinese:
                # 使用 PyMuPDF 内置的 china-ss 字体（唯一可靠支持中文的方案）
                try:
                    page.insert_text(
                        (text_x, y),
                        page_text,
                        fontsize=font_size,
                        color=color_normalized,
                        fontname="china-ss"
                    )
                    text_inserted = True
                    logger.info(f"成功使用china-ss字体插入页码: {page_text}, 位置: ({text_x}, {y})")
                except Exception as e:
                    logger.warning(f"使用china-ss字体失败: {str(e)}")
            
            # 如果不包含中文或china-ss失败，使用默认字体
            if not text_inserted:
                try:
                    page.insert_text(
                        (text_x, y),
                        page_text,
                        fontsize=font_size,
                        color=color_normalized
                    )
                    text_inserted = True
                    logger.info(f"成功使用默认字体插入页码: {page_text}, 位置: ({text_x}, {y})")
                except Exception as e:
                    logger.error(f"插入页码文本失败: {str(e)}, 位置: ({text_x}, {y}), 文本: {page_text}")
                    # 不抛出异常，继续处理其他页面
        
        # 保存结果
        output_filename = f"{timestamp}_pagenum_{original_filename}"
        output_path = os.path.join(app.config['CONVERTED_FOLDER'], output_filename)
        doc.save(output_path, garbage=4, deflate=True)  # 优化保存
        doc.close()
        
        # 删除临时文件
        try:
            os.remove(temp_input_path)
        except:
            pass
        
        elapsed_time = time.time() - start_time
        file_size_mb = os.path.getsize(output_path) / (1024 * 1024)
        
        result = {
            'success': True,
            'message': '页码添加成功',
            'filename': output_filename,
            'url': f'/download/{output_filename}',
            'total_pages': total_pages,
            'pages_with_numbers': total_pages - start_page + 1,
            'file_size': os.path.getsize(output_path),
            'file_size_mb': round(file_size_mb, 2),
            'elapsed_time': round(elapsed_time, 2)
        }
        
        logger.info(f"页码添加成功: {output_filename}, 耗时{elapsed_time:.2f}秒, 大小{file_size_mb:.2f}MB")
        return jsonify(result)
        
    except Exception as e:
        logger.error(f"添加页码失败: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': f'添加页码失败: {str(e)}'}), 500


def to_roman(n):
    """将数字转换为罗马数字"""
    val = [
        1000, 900, 500, 400,
        100, 90, 50, 40,
        10, 9, 5, 4,
        1
    ]
    syb = [
        "M", "CM", "D", "CD",
        "C", "XC", "L", "XL",
        "X", "IX", "V", "IV",
        "I"
    ]
    roman_num = ''
    i = 0
    while n > 0:
        for _ in range(n // val[i]):
            roman_num += syb[i]
            n -= val[i]
        i += 1
    return roman_num


def to_chinese_number(n):
    """将数字转换为中文数字（第一页、第二页等）"""
    chinese_digits = ['', '一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
    
    if n <= 10:
        return '第' + chinese_digits[n] + '页'
    elif n <= 99:
        tens = n // 10
        ones = n % 10
        if tens == 1:
            if ones == 0:
                return '第十页'
            else:
                return '第十' + chinese_digits[ones] + '页'
        else:
            if ones == 0:
                return '第' + chinese_digits[tens] + '十页'
            else:
                return '第' + chinese_digits[tens] + '十' + chinese_digits[ones] + '页'
    else:
        # 超过99页，使用数字
        return '第' + str(n) + '页'


def format_page_number(page_num, total_pages, style):
    """
    根据样式格式化页码文本
    
    样式选项：
    1. chinese_first: 第1页
    2. chinese_fraction: 第1/20页
    3. chinese_total: 第1页,共20页
    4. simple: 1
    5. dash: -1-
    6. fraction: 1/20
    7. page_en: Page 1
    8. page_dot: P.1
    """
    if style == 'chinese_first':
        return f"第{page_num}页"
    elif style == 'chinese_fraction':
        return f"第{page_num}/{total_pages}页"
    elif style == 'chinese_total':
        return f"第{page_num}页,共{total_pages}页"
    elif style == 'simple':
        return str(page_num)
    elif style == 'dash':
        return f"-{page_num}-"
    elif style == 'fraction':
        return f"{page_num}/{total_pages}"
    elif style == 'page_en':
        return f"Page {page_num}"
    elif style == 'page_dot':
        return f"P.{page_num}"
    else:
        # 默认返回简单数字
        return str(page_num)


@app.route('/pdf/preview-page-number', methods=['POST'])
def preview_page_number():
    """
    预览页码效果
    返回第一页添加页码后的预览图
    """
    try:
        # 检查是否有文件
        if 'file' not in request.files:
            return jsonify({'error': '请选择PDF文件'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': '请选择PDF文件'}), 400
        
        # 获取参数
        position = request.form.get('position', 'bottom_center')
        style = request.form.get('style', 'simple')
        font_size = int(request.form.get('font_size', '12'))
        margin = int(request.form.get('margin', '20'))
        color = request.form.get('color', '#000000')
        start_page = int(request.form.get('start_page', '1'))
        preview_page = int(request.form.get('preview_page', '1'))  # 指定预览的页码
        
        # 解析颜色
        try:
            color_rgb = tuple(int(color[i:i+2], 16) for i in (1, 3, 5))
            color_normalized = tuple(c / 255.0 for c in color_rgb)
        except:
            color_normalized = (0, 0, 0)
        
        # 保存临时文件
        timestamp = int(time.time() * 1000)
        temp_path = os.path.join(app.config['UPLOAD_FOLDER'], f"temp_preview_{timestamp}.pdf")
        file.save(temp_path)
        
        # 打开PDF
        doc = fitz.open(temp_path)
        if len(doc) == 0:
            doc.close()
            os.remove(temp_path)
            return jsonify({'error': 'PDF文件为空'}), 400
        
        total_pages = len(doc)
        
        # 验证预览页码
        if preview_page < 1 or preview_page > total_pages:
            preview_page = 1
        
        # 处理指定页面（预览）
        page = doc[preview_page - 1]  # 转换为0-based索引
        rect = page.rect
        
        # 计算页码文本（预览时总是显示页码，方便预览效果）
        # 如果预览页小于起始页，显示预览页的实际页码；否则显示从起始页开始的页码
        if preview_page >= start_page:
            # 计算实际页码（从start_page开始计数，从1开始）
            actual_page_num = preview_page - start_page + 1
            # 计算实际需要添加页码的总页数（用于样式中的总页数显示）
            actual_total_pages = total_pages - start_page + 1
            page_text = format_page_number(actual_page_num, actual_total_pages, style)
        else:
            # 预览页小于起始页，显示预览页的实际页码（用于预览效果）
            page_text = format_page_number(preview_page, total_pages, style)
        
        # 总是添加页码（预览时）
        if True:  # 预览时总是显示页码
            
            # 计算位置和对齐方式
            if position == 'top_left':
                x = margin
                y = margin
                align = 0  # 左对齐
            elif position == 'top_center':
                x = rect.width / 2
                y = margin
                align = 1  # 居中
            elif position == 'top_right':
                x = rect.width - margin
                y = margin
                align = 2  # 右对齐
            elif position == 'bottom_left':
                x = margin
                y = rect.height - margin
                align = 0
            elif position == 'bottom_center':
                x = rect.width / 2
                y = rect.height - margin
                align = 1
            elif position == 'bottom_right':
                x = rect.width - margin
                y = rect.height - margin
                align = 2
            else:
                x = rect.width / 2
                y = rect.height - margin
                align = 1
            
            # 使用insert_text，根据对齐方式手动计算文本位置
            # 更精确地估算文本宽度（区分中文字符和数字/字母）
            # 对于数字和字母，平均宽度约为字体大小的0.5-0.6倍
            # 对于中文，平均宽度约为字体大小
            text_width = 0
            for char in page_text:
                # 判断是否为中文字符（包括中文标点）
                if '\u4e00' <= char <= '\u9fff' or char in '，。、；：？！':
                    text_width += font_size  # 中文字符宽度约为字体大小
                else:
                    text_width += font_size * 0.55  # 数字和字母宽度约为字体大小的0.55倍
            text_height = font_size * 1.2  # 文本高度
            
            # 根据对齐方式调整x坐标
            # 增加一些宽度余量，确保文本完全显示（避免省略号）
            text_width_with_margin = text_width * 1.1  # 增加10%的余量
            
            if align == 1:  # 居中对齐
                # x已经是中心点，需要减去文本宽度的一半
                text_x = x - text_width_with_margin / 2
            elif align == 2:  # 右对齐
                # x是右边界，需要减去文本宽度
                text_x = x - text_width_with_margin
            else:  # 左对齐 (align == 0)
                # x就是左边界，直接使用
                text_x = x
            
            # 确保文本不超出页面边界
            if text_x < 0:
                text_x = 0
            elif text_x + text_width_with_margin > rect.width:
                text_x = rect.width - text_width_with_margin
                if text_x < 0:
                    text_x = 0
            
            # 添加页码文本
            # 检查文本是否包含中文字符，如果包含则使用支持中文的字体
            has_chinese = any('\u4e00' <= char <= '\u9fff' for char in page_text)
            
            text_inserted = False
            
            if has_chinese:
                # 使用 PyMuPDF 内置的 china-ss 字体（唯一可靠支持中文的方案）
                try:
                    page.insert_text(
                        (text_x, y),
                        page_text,
                        fontsize=font_size,
                        color=color_normalized,
                        fontname="china-ss"
                    )
                    text_inserted = True
                    logger.info(f"成功使用china-ss字体插入页码: {page_text}, 位置: ({text_x}, {y})")
                except Exception as e:
                    logger.warning(f"使用china-ss字体失败: {str(e)}")
            
            # 如果不包含中文或china-ss失败，使用默认字体
            if not text_inserted:
                try:
                    page.insert_text(
                        (text_x, y),
                        page_text,
                        fontsize=font_size,
                        color=color_normalized
                    )
                    text_inserted = True
                    logger.info(f"成功使用默认字体插入页码: {page_text}, 位置: ({text_x}, {y})")
                except Exception as e:
                    logger.error(f"插入页码文本失败: {str(e)}, 位置: ({text_x}, {y}), 文本: {page_text}")
                    # 不抛出异常，继续处理其他页面
            
            # 在页码周围绘制红色矩形框（只包住页码，不包整个页面）
            # 矩形框内边距（小一点，更紧凑）
            padding = 3  # 矩形框内边距
            
            # 计算页码文本的精确边界框
            # PDF坐标系：y=0在底部，文本的y坐标是基线位置
            # 文本从基线向上延伸，所以矩形框的y0应该在y-text_height，y1在y附近
            rect_x0 = text_x - padding
            rect_y0 = y - text_height - padding  # 文本顶部
            rect_x1 = text_x + text_width + padding
            rect_y1 = y + padding  # 文本底部（基线稍微向下一点）
            
            # 确保矩形框坐标有效（不超出页面边界）
            rect_x0 = max(0, rect_x0)
            rect_y0 = max(0, rect_y0)
            rect_x1 = min(rect.width, rect_x1)
            rect_y1 = min(rect.height, rect_y1)
            
            # 绘制红色矩形框（只包住页码）
            page_num_rect = fitz.Rect(rect_x0, rect_y0, rect_x1, rect_y1)
            page.draw_rect(page_num_rect, color=(1, 0, 0), width=2)  # 红色，线宽2
        
        # 渲染为图片（缩略图）
        zoom = 1.0
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        
        # 转换为base64
        img_data = pix.tobytes("png")
        img_base64 = base64.b64encode(img_data).decode('utf-8')
        
        doc.close()
        os.remove(temp_path)
        
        return jsonify({
            'success': True,
            'preview': f'data:image/png;base64,{img_base64}',
            'total_pages': total_pages
        })
        
    except Exception as e:
        logger.error(f"预览页码失败: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': f'预览失败: {str(e)}'}), 500


@app.errorhandler(413)
def request_entity_too_large(error):
    """文件过大错误处理"""
    return jsonify({'error': '文件大小超过50MB限制'}), 413


@app.errorhandler(500)
def internal_server_error(error):
    """内部服务器错误处理"""
    logger.error(f"服务器错误: {str(error)}")
    return jsonify({'error': '服务器内部错误'}), 500


@app.route('/pdf/crop-preview', methods=['POST'])
def crop_preview():
    """
    生成PDF裁剪预览图
    返回第一页的预览图，显示裁剪区域
    """
    try:
        if 'file' not in request.files:
            return jsonify({'error': '请选择PDF文件'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': '请选择PDF文件'}), 400
        
        # 获取参数
        crop_type = request.form.get('crop_type', 'none')  # none, left_right, top_bottom, custom
        crop_box = request.form.get('crop_box', None)  # JSON格式: {"x0": 0, "y0": 0, "x1": 100, "y1": 100}
        preview_page = int(request.form.get('preview_page', '1'))
        
        # 保存临时文件
        timestamp = int(time.time() * 1000)
        temp_path = os.path.join(app.config['UPLOAD_FOLDER'], f"temp_crop_preview_{timestamp}.pdf")
        file.save(temp_path)
        
        try:
            # 打开PDF
            doc = fitz.open(temp_path)
            if len(doc) == 0:
                doc.close()
                os.remove(temp_path)
                return jsonify({'error': 'PDF文件为空'}), 400
            
            total_pages = len(doc)
            if preview_page < 1 or preview_page > total_pages:
                preview_page = 1
            
            page = doc[preview_page - 1]
            rect = page.rect
            
            # 生成纯净的预览图（不绘制任何裁剪框）
            # 裁剪框和分隔线将在前端通过 CSS 实现
            preview_page_obj = doc[preview_page - 1]
            
            # 生成预览图
            mat = fitz.Matrix(2, 2)  # 2倍缩放，提高清晰度
            pix = preview_page_obj.get_pixmap(matrix=mat, alpha=False)
            img_data = pix.tobytes("png")
            img_base64 = base64.b64encode(img_data).decode('utf-8')
            
            doc.close()
            os.remove(temp_path)
            
            return jsonify({
                'success': True,
                'preview': f'data:image/png;base64,{img_base64}',
                'page_width': int(rect.width),
                'page_height': int(rect.height),
                'total_pages': total_pages
            })
            
        except Exception as e:
            logger.error(f"生成裁剪预览失败: {str(e)}")
            logger.error(traceback.format_exc())
            if os.path.exists(temp_path):
                os.remove(temp_path)
            return jsonify({'error': f'生成预览失败: {str(e)}'}), 500
            
    except Exception as e:
        logger.error(f"裁剪预览错误: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': f'处理失败: {str(e)}'}), 500


@app.route('/pdf/crop', methods=['POST'])
def crop_pdf():
    """
    PDF页面裁剪
    支持左右分割、上下分割、自定义裁剪
    对所有页面执行一致操作
    """
    try:
        if 'file' not in request.files:
            return jsonify({'error': '请选择PDF文件'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': '请选择PDF文件'}), 400
        
        # 获取参数
        crop_type = request.form.get('crop_type', 'none')  # none, left_right, top_bottom, custom
        crop_box = request.form.get('crop_box', None)  # JSON格式: {"x0": 0, "y0": 0, "x1": 100, "y1": 100}
        export_size = request.form.get('export_size', 'original')  # original, a4
        
        # 保存上传的文件
        timestamp = int(time.time() * 1000)
        original_filename = secure_filename(file.filename)
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{timestamp}_crop_{original_filename}")
        file.save(input_path)
        
        try:
            start_time = time.time()
            
            # 打开PDF
            doc = fitz.open(input_path)
            if len(doc) == 0:
                doc.close()
                os.remove(input_path)
                return jsonify({'error': 'PDF文件为空'}), 400
            
            total_pages = len(doc)
            logger.info(f"开始裁剪PDF: {original_filename}, 总页数: {total_pages}, 裁剪类型: {crop_type}")
            
            # 创建新PDF
            new_doc = fitz.open()
            
            # A4尺寸（点，1点=1/72英寸）
            A4_WIDTH = 595.276  # A4宽度
            A4_HEIGHT = 841.890  # A4高度
            
            # 解析裁剪框坐标
            base_crop_rect = None
            if crop_box:
                try:
                    import json
                    box = json.loads(crop_box)
                    base_crop_rect = fitz.Rect(box['x0'], box['y0'], box['x1'], box['y1'])
                    logger.info(f"使用自定义裁剪框: {base_crop_rect}")
                except Exception as e:
                    logger.warning(f"裁剪框参数解析失败: {str(e)}")
            
            # 处理每一页
            for page_num in range(total_pages):
                page = doc[page_num]
                rect = page.rect
                
                # 确定基础裁剪区域（如果有裁剪框则使用裁剪框，否则使用整页）
                if base_crop_rect:
                    crop_area = base_crop_rect
                else:
                    crop_area = rect
                
                if crop_type == 'left_right':
                    # 左右分割：在裁剪框内从中间分成两页
                    left_rect = fitz.Rect(crop_area.x0, crop_area.y0, crop_area.x0 + crop_area.width / 2, crop_area.y1)
                    right_rect = fitz.Rect(crop_area.x0 + crop_area.width / 2, crop_area.y0, crop_area.x1, crop_area.y1)
                    
                    # 创建左半页
                    if export_size == 'a4':
                        new_page = new_doc.new_page(width=A4_WIDTH, height=A4_HEIGHT)
                        # 计算缩放比例以适应A4
                        scale_x = A4_WIDTH / left_rect.width
                        scale_y = A4_HEIGHT / left_rect.height
                        scale = min(scale_x, scale_y)
                        new_rect = fitz.Rect(0, 0, left_rect.width * scale, left_rect.height * scale)
                        new_rect = new_rect.center_in(new_page.rect)
                        new_page.show_pdf_page(new_rect, doc, page_num, clip=left_rect)
                    else:
                        new_page = new_doc.new_page(width=left_rect.width, height=left_rect.height)
                        new_page.show_pdf_page(new_page.rect, doc, page_num, clip=left_rect)
                    
                    # 创建右半页
                    if export_size == 'a4':
                        new_page = new_doc.new_page(width=A4_WIDTH, height=A4_HEIGHT)
                        scale_x = A4_WIDTH / right_rect.width
                        scale_y = A4_HEIGHT / right_rect.height
                        scale = min(scale_x, scale_y)
                        new_rect = fitz.Rect(0, 0, right_rect.width * scale, right_rect.height * scale)
                        new_rect = new_rect.center_in(new_page.rect)
                        new_page.show_pdf_page(new_rect, doc, page_num, clip=right_rect)
                    else:
                        new_page = new_doc.new_page(width=right_rect.width, height=right_rect.height)
                        new_page.show_pdf_page(new_page.rect, doc, page_num, clip=right_rect)
                    
                elif crop_type == 'top_bottom':
                    # 上下分割：在裁剪框内从中间分成两页
                    top_rect = fitz.Rect(crop_area.x0, crop_area.y0, crop_area.x1, crop_area.y0 + crop_area.height / 2)
                    bottom_rect = fitz.Rect(crop_area.x0, crop_area.y0 + crop_area.height / 2, crop_area.x1, crop_area.y1)
                    
                    # 创建上半页
                    if export_size == 'a4':
                        new_page = new_doc.new_page(width=A4_WIDTH, height=A4_HEIGHT)
                        scale_x = A4_WIDTH / top_rect.width
                        scale_y = A4_HEIGHT / top_rect.height
                        scale = min(scale_x, scale_y)
                        new_rect = fitz.Rect(0, 0, top_rect.width * scale, top_rect.height * scale)
                        new_rect = new_rect.center_in(new_page.rect)
                        new_page.show_pdf_page(new_rect, doc, page_num, clip=top_rect)
                    else:
                        new_page = new_doc.new_page(width=top_rect.width, height=top_rect.height)
                        new_page.show_pdf_page(new_page.rect, doc, page_num, clip=top_rect)
                    
                    # 创建下半页
                    if export_size == 'a4':
                        new_page = new_doc.new_page(width=A4_WIDTH, height=A4_HEIGHT)
                        scale_x = A4_WIDTH / bottom_rect.width
                        scale_y = A4_HEIGHT / bottom_rect.height
                        scale = min(scale_x, scale_y)
                        new_rect = fitz.Rect(0, 0, bottom_rect.width * scale, bottom_rect.height * scale)
                        new_rect = new_rect.center_in(new_page.rect)
                        new_page.show_pdf_page(new_rect, doc, page_num, clip=bottom_rect)
                    else:
                        new_page = new_doc.new_page(width=bottom_rect.width, height=bottom_rect.height)
                        new_page.show_pdf_page(new_page.rect, doc, page_num, clip=bottom_rect)
                    
                elif crop_type == 'custom' or crop_type == 'none':
                    # 自定义裁剪或不分割：直接使用裁剪框区域
                    if export_size == 'a4':
                        new_page = new_doc.new_page(width=A4_WIDTH, height=A4_HEIGHT)
                        scale_x = A4_WIDTH / crop_area.width
                        scale_y = A4_HEIGHT / crop_area.height
                        scale = min(scale_x, scale_y)
                        new_rect = fitz.Rect(0, 0, crop_area.width * scale, crop_area.height * scale)
                        new_rect = new_rect.center_in(new_page.rect)
                        new_page.show_pdf_page(new_rect, doc, page_num, clip=crop_area)
                    else:
                        new_page = new_doc.new_page(width=crop_area.width, height=crop_area.height)
                        new_page.show_pdf_page(new_page.rect, doc, page_num, clip=crop_area)
            
            # 保存结果
            output_filename = f"{timestamp}_crop_{original_filename}"
            output_path = os.path.join(app.config['CONVERTED_FOLDER'], output_filename)
            new_doc.save(output_path, garbage=4, deflate=True)
            
            # 获取文件大小
            file_size = os.path.getsize(output_path)
            file_size_mb = file_size / (1024 * 1024)
            
            elapsed_time = time.time() - start_time
            new_total_pages = len(new_doc)
            
            new_doc.close()
            doc.close()
            os.remove(input_path)
            
            logger.info(f"PDF裁剪成功: {output_filename}, 耗时{elapsed_time:.2f}秒, 大小{file_size_mb:.2f}MB, 原页数{total_pages}, 新页数{new_total_pages}")
            
            return jsonify({
                'success': True,
                'filename': output_filename,
                'original_pages': total_pages,
                'new_pages': new_total_pages,
                'file_size': file_size,
                'file_size_mb': round(file_size_mb, 2),
                'elapsed_time': round(elapsed_time, 2)
            })
            
        except Exception as e:
            logger.error(f"PDF裁剪失败: {str(e)}")
            logger.error(traceback.format_exc())
            if os.path.exists(input_path):
                os.remove(input_path)
            return jsonify({'error': f'裁剪失败: {str(e)}'}), 500
            
    except Exception as e:
        logger.error(f"PDF裁剪错误: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': f'处理失败: {str(e)}'}), 500


@app.route('/pdf/combine-preview', methods=['POST'])
def combine_preview():
    """
    生成页面合并预览
    """
    try:
        if 'file' not in request.files:
            return jsonify({'error': '请选择PDF文件'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': '请选择PDF文件'}), 400
        
        # 获取参数
        merge_type = request.form.get('merge_type', 'left_right')  # left_right, top_bottom
        pair_index = request.form.get('pair_index', None)  # 指定页面对索引（按需加载）
        left_page = request.form.get('left_page', None)  # 左边页
        right_page = request.form.get('right_page', None)  # 右边页
        
        # 保存上传的文件
        timestamp = int(time.time() * 1000)
        original_filename = secure_filename(file.filename)
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{timestamp}_combine_preview_{original_filename}")
        file.save(input_path)
        
        try:
            # 打开PDF
            doc = fitz.open(input_path)
            total_pages = len(doc)
            
            if total_pages == 0:
                doc.close()
                os.remove(input_path)
                return jsonify({'error': 'PDF文件为空'}), 400
            
            # 生成页面缩略图预览
            previews = []
            
            # 如果指定了页面对，只生成该页面对的预览
            if pair_index is not None and left_page is not None:
                left_page_num = int(left_page)
                right_page_num = int(right_page) if right_page and right_page != '-1' else None
                
                # 创建合并预览
                left_page_obj = doc[left_page_num]
                left_pix = left_page_obj.get_pixmap(matrix=fitz.Matrix(0.5, 0.5))
                
                if right_page_num is not None and right_page_num < total_pages:
                    right_page_obj = doc[right_page_num]
                    right_pix = right_page_obj.get_pixmap(matrix=fitz.Matrix(0.5, 0.5))
                    
                    # 合并两个缩略图
                    from PIL import Image
                    import io
                    
                    left_img = Image.frombytes("RGB", [left_pix.width, left_pix.height], left_pix.samples)
                    right_img = Image.frombytes("RGB", [right_pix.width, right_pix.height], right_pix.samples)
                    
                    if merge_type == 'left_right':
                        # 左右合并
                        combined_width = left_pix.width + right_pix.width
                        combined_height = max(left_pix.height, right_pix.height)
                        combined_img = Image.new('RGB', (combined_width, combined_height), 'white')
                        combined_img.paste(left_img, (0, 0))
                        combined_img.paste(right_img, (left_pix.width, 0))
                    else:
                        # 上下合并
                        combined_width = max(left_pix.width, right_pix.width)
                        combined_height = left_pix.height + right_pix.height
                        combined_img = Image.new('RGB', (combined_width, combined_height), 'white')
                        combined_img.paste(left_img, (0, 0))
                        combined_img.paste(right_img, (0, left_pix.height))
                    
                    # 转换为base64
                    buffer = io.BytesIO()
                    combined_img.save(buffer, format='JPEG', quality=85)
                    img_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
                else:
                    # 只有一页，直接使用
                    img_base64 = base64.b64encode(left_pix.tobytes('jpeg')).decode('utf-8')
                
                previews.append({
                    'pair_index': int(pair_index),
                    'left_page': left_page_num,
                    'right_page': right_page_num,
                    'image': f'data:image/jpeg;base64,{img_base64}'
                })
                
            else:
                # 生成所有页面对的预览（兼容旧版本）
                page_index = 0
                
                while page_index < total_pages:
                    # 计算当前页对
                    left_page_num = page_index
                    right_page_num = page_index + 1 if page_index + 1 < total_pages else None
                    
                    # 创建合并预览
                    left_page_obj = doc[left_page_num]
                    left_pix = left_page_obj.get_pixmap(matrix=fitz.Matrix(0.3, 0.3))
                    
                    if right_page_num is not None:
                        right_page_obj = doc[right_page_num]
                        right_pix = right_page_obj.get_pixmap(matrix=fitz.Matrix(0.3, 0.3))
                        
                        # 合并两个缩略图
                        if merge_type == 'left_right':
                            # 左右合并
                            combined_width = left_pix.width + right_pix.width
                            combined_height = max(left_pix.height, right_pix.height)
                        else:
                            # 上下合并
                            combined_width = max(left_pix.width, right_pix.width)
                            combined_height = left_pix.height + right_pix.height
                        
                        # 创建新的pixmap
                        from PIL import Image
                        import io
                        
                        left_img = Image.frombytes("RGB", [left_pix.width, left_pix.height], left_pix.samples)
                        right_img = Image.frombytes("RGB", [right_pix.width, right_pix.height], right_pix.samples)
                        
                        if merge_type == 'left_right':
                            combined_img = Image.new('RGB', (combined_width, combined_height), 'white')
                            combined_img.paste(left_img, (0, 0))
                            combined_img.paste(right_img, (left_pix.width, 0))
                        else:
                            combined_img = Image.new('RGB', (combined_width, combined_height), 'white')
                            combined_img.paste(left_img, (0, 0))
                            combined_img.paste(right_img, (0, left_pix.height))
                        
                        # 转换为base64
                        buffer = io.BytesIO()
                        combined_img.save(buffer, format='JPEG', quality=85)
                        img_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
                    else:
                        # 只有一页，直接使用
                        img_base64 = base64.b64encode(left_pix.tobytes('jpeg')).decode('utf-8')
                    
                    previews.append({
                        'pair_index': len(previews),
                        'left_page': left_page_num,
                        'right_page': right_page_num,
                        'image': f'data:image/jpeg;base64,{img_base64}'
                    })
                    
                    page_index += 2
            
            doc.close()
            os.remove(input_path)
            
            return jsonify({
                'success': True,
                'total_pages': total_pages,
                'previews': previews
            })
            
        except Exception as e:
            logger.error(f"生成合并预览失败: {str(e)}")
            logger.error(traceback.format_exc())
            if os.path.exists(input_path):
                os.remove(input_path)
            return jsonify({'error': f'生成预览失败: {str(e)}'}), 500
            
    except Exception as e:
        logger.error(f"合并预览错误: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': f'处理失败: {str(e)}'}), 500


@app.route('/pdf/combine-pages', methods=['POST'])
def combine_pages():
    """
    PDF页面合并
    支持左右合并、上下合并
    可设置独立页、旋转页面
    """
    try:
        if 'file' not in request.files:
            return jsonify({'error': '请选择PDF文件'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': '请选择PDF文件'}), 400
        
        # 获取参数
        merge_type = request.form.get('merge_type', 'left_right')  # left_right, top_bottom
        merge_plan_json = request.form.get('merge_plan', '[]')  # JSON数组
        
        # 解析合并计划
        import json
        merge_plan = json.loads(merge_plan_json)
        
        # 保存上传的文件
        timestamp = int(time.time() * 1000)
        original_filename = secure_filename(file.filename)
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{timestamp}_combine_{original_filename}")
        file.save(input_path)
        
        try:
            start_time = time.time()
            
            # 打开PDF
            doc = fitz.open(input_path)
            total_pages = len(doc)
            
            if total_pages == 0:
                doc.close()
                os.remove(input_path)
                return jsonify({'error': 'PDF文件为空'}), 400
            
            logger.info(f"开始合并PDF: {original_filename}, 总页数: {total_pages}, 合并类型: {merge_type}")
            
            # 创建新PDF
            new_doc = fitz.open()
            
            # 根据合并计划处理页面
            for plan_item in merge_plan:
                left_page_num = plan_item.get('left_page')
                right_page_num = plan_item.get('right_page')
                is_independent = plan_item.get('is_independent', False)
                left_rotation = plan_item.get('left_rotation', 0)
                right_rotation = plan_item.get('right_rotation', 0)
                
                if is_independent:
                    # 设置为独立页，分别添加
                    if left_page_num is not None:
                        page = doc[left_page_num]
                        new_page = new_doc.new_page(width=page.rect.width, height=page.rect.height)
                        new_page.show_pdf_page(new_page.rect, doc, left_page_num, rotate=left_rotation)
                    
                    if right_page_num is not None:
                        page = doc[right_page_num]
                        new_page = new_doc.new_page(width=page.rect.width, height=page.rect.height)
                        new_page.show_pdf_page(new_page.rect, doc, right_page_num, rotate=right_rotation)
                else:
                    # 合并页面
                    if left_page_num is not None and right_page_num is not None:
                        left_page = doc[left_page_num]
                        right_page = doc[right_page_num]
                        
                        # 获取旋转后的页面尺寸
                        left_rect = left_page.rect
                        right_rect = right_page.rect
                        
                        if merge_type == 'left_right':
                            # 左右合并
                            new_width = left_rect.width + right_rect.width
                            new_height = max(left_rect.height, right_rect.height)
                            new_page = new_doc.new_page(width=new_width, height=new_height)
                            
                            # 放置左页
                            left_target = fitz.Rect(0, 0, left_rect.width, left_rect.height)
                            new_page.show_pdf_page(left_target, doc, left_page_num, rotate=left_rotation)
                            
                            # 放置右页
                            right_target = fitz.Rect(left_rect.width, 0, new_width, right_rect.height)
                            new_page.show_pdf_page(right_target, doc, right_page_num, rotate=right_rotation)
                        else:
                            # 上下合并
                            new_width = max(left_rect.width, right_rect.width)
                            new_height = left_rect.height + right_rect.height
                            new_page = new_doc.new_page(width=new_width, height=new_height)
                            
                            # 放置上页（左页）
                            top_target = fitz.Rect(0, 0, left_rect.width, left_rect.height)
                            new_page.show_pdf_page(top_target, doc, left_page_num, rotate=left_rotation)
                            
                            # 放置下页（右页）
                            bottom_target = fitz.Rect(0, left_rect.height, right_rect.width, new_height)
                            new_page.show_pdf_page(bottom_target, doc, right_page_num, rotate=right_rotation)
                    elif left_page_num is not None:
                        # 只有左页
                        page = doc[left_page_num]
                        new_page = new_doc.new_page(width=page.rect.width, height=page.rect.height)
                        new_page.show_pdf_page(new_page.rect, doc, left_page_num, rotate=left_rotation)
            
            # 保存PDF
            output_filename = f"{timestamp}_combined_{original_filename}"
            output_path = os.path.join(app.config['CONVERTED_FOLDER'], output_filename)
            new_doc.save(output_path, garbage=4, deflate=True)
            
            # 获取文件大小
            file_size = os.path.getsize(output_path)
            file_size_mb = file_size / (1024 * 1024)
            
            elapsed_time = time.time() - start_time
            new_total_pages = len(new_doc)
            
            new_doc.close()
            doc.close()
            os.remove(input_path)
            
            logger.info(f"PDF合并成功: {output_filename}, 耗时{elapsed_time:.2f}秒, 大小{file_size_mb:.2f}MB, 原页数{total_pages}, 新页数{new_total_pages}")
            
            return jsonify({
                'success': True,
                'filename': output_filename,
                'original_pages': total_pages,
                'new_pages': new_total_pages,
                'file_size': file_size,
                'file_size_mb': round(file_size_mb, 2),
                'elapsed_time': round(elapsed_time, 2)
            })
            
        except Exception as e:
            logger.error(f"PDF合并失败: {str(e)}")
            logger.error(traceback.format_exc())
            if os.path.exists(input_path):
                os.remove(input_path)
            return jsonify({'error': f'合并失败: {str(e)}'}), 500
            
    except Exception as e:
        logger.error(f"PDF合并错误: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': f'处理失败: {str(e)}'}), 500


@app.route('/pdf/watermark-password', methods=['POST'])
def watermark_password():
    """
    PDF加水印/密码 - 统一接口
    支持：仅水印、仅密码、水印+密码
    限制：水印30页，文件50MB
    """
    try:
        if 'file' not in request.files:
            return jsonify({'error': '请选择PDF文件'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': '请选择PDF文件'}), 400
        
        # 获取功能选择
        add_watermark = request.form.get('add_watermark', 'false') == 'true'
        add_password = request.form.get('add_password', 'false') == 'true'
        
        if not add_watermark and not add_password:
            return jsonify({'error': '请至少选择一个功能'}), 400
        
        # 水印参数
        watermark_text = request.form.get('watermark_text', '水印')
        position = request.form.get('position', 'center')
        
        try:
            fontsize = int(request.form.get('fontsize', '50'))
            color_r = float(request.form.get('color_r', '0.8'))
            color_g = float(request.form.get('color_g', '0.8'))
            color_b = float(request.form.get('color_b', '0.8'))
            rotation_str = request.form.get('rotation', '45')
            rotation = int(float(rotation_str)) if rotation_str else 45
            opacity = float(request.form.get('opacity', '0.3'))  # 透明度 0-1
            font_style = request.form.get('font_style', 'song')  # 字体样式
        except ValueError as e:
            logger.error(f"参数转换错误: {e}")
            return jsonify({'error': '参数格式错误'}), 400
        
        # 密码参数
        user_pw = request.form.get('user_password', '')
        owner_pw = request.form.get('owner_password', '')
        allow_print = request.form.get('allow_print', 'true') == 'true'
        allow_copy = request.form.get('allow_copy', 'true') == 'true'
        allow_modify = request.form.get('allow_modify', 'false') == 'true'
        allow_annotate = request.form.get('allow_annotate', 'false') == 'true'
        
        if add_password and not user_pw:
            return jsonify({'error': '请输入用户密码'}), 400
        
        if not owner_pw:
            owner_pw = user_pw
        
        # 保存文件
        timestamp = int(time.time() * 1000)
        original_filename = secure_filename(file.filename)
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{timestamp}_{original_filename}")
        file.save(input_path)
        
        try:
            start_time = time.time()
            file_size = os.path.getsize(input_path)
            file_size_mb = file_size / (1024 * 1024)
            
            if file_size_mb > 50:
                os.remove(input_path)
                return jsonify({'error': '文件大小超过50MB限制'}), 400
            
            doc = fitz.open(input_path)
            total_pages = len(doc)
            
            if total_pages == 0:
                doc.close()
                os.remove(input_path)
                return jsonify({'error': 'PDF文件为空'}), 400
            
            if add_watermark and total_pages > 30:
                doc.close()
                os.remove(input_path)
                return jsonify({'error': '水印功能仅支持30页以内'}), 400
            
            logger.info(f"处理PDF: {original_filename}, 页数: {total_pages}, 水印: {add_watermark}, 密码: {add_password}")
            
            # 添加水印
            if add_watermark and watermark_text:
                import math
                
                # 字体映射
                font_map = {
                    'song': 'china-s',      # 宋体
                    'hei': 'china-ss',      # 黑体
                    'kai': 'china-s',       # 楷体（使用宋体代替）
                    'fangsong': 'china-s',  # 仿宋（使用宋体代替）
                }
                fontname = font_map.get(font_style, 'china-s')
                
                for page_num in range(total_pages):
                    page = doc[page_num]
                    rect = page.rect
                    
                    if position == 'tile':
                        # 平铺水印模式（支持任意角度）
                        text_width = len(watermark_text) * fontsize * 0.7
                        text_height = fontsize
                        
                        # 水印间距
                        spacing_x = text_width + fontsize * 3
                        spacing_y = fontsize * 4
                        
                        # 扩大绘制范围以覆盖旋转后的区域
                        margin = max(rect.width, rect.height)
                        
                        # 使用TextWriter实现任意角度旋转
                        y_pos = -margin
                        row = 0
                        while y_pos < rect.height + margin:
                            x_offset = (spacing_x / 2) if row % 2 else 0
                            x_pos = -margin + x_offset
                            
                            while x_pos < rect.width + margin:
                                # 创建旋转矩阵
                                angle_rad = math.radians(-rotation)  # 负号使文字逆时针旋转
                                cos_a, sin_a = math.cos(angle_rad), math.sin(angle_rad)
                                
                                # 计算旋转后的位置
                                rot_matrix = fitz.Matrix(cos_a, sin_a, -sin_a, cos_a, x_pos, y_pos)
                                
                                # 使用Shape绘制带旋转的文本
                                tw = fitz.TextWriter(page.rect)
                                tw.append((0, fontsize), watermark_text, fontsize=fontsize, font=fitz.Font(fontname))
                                tw.write_text(page, morph=(fitz.Point(0, 0), rot_matrix), color=(color_r, color_g, color_b), opacity=opacity)
                                
                                x_pos += spacing_x
                            
                            y_pos += spacing_y
                            row += 1
                    else:
                        # 单点水印模式（支持任意角度）
                        # 计算位置
                        if position == 'center':
                            x, y = rect.width / 2, rect.height / 2
                        elif position == 'top_left':
                            x, y = rect.width * 0.15, rect.height * 0.15
                        elif position == 'top_right':
                            x, y = rect.width * 0.85, rect.height * 0.15
                        elif position == 'bottom_left':
                            x, y = rect.width * 0.15, rect.height * 0.85
                        elif position == 'bottom_right':
                            x, y = rect.width * 0.85, rect.height * 0.85
                        else:
                            x, y = rect.width / 2, rect.height / 2
                        
                        # 应用透明度：通过颜色混合模拟透明度
                        if opacity < 1.0:
                            # 将颜色与白色背景混合来模拟透明度
                            blended_r = color_r * opacity + (1 - opacity) * 1.0
                            blended_g = color_g * opacity + (1 - opacity) * 1.0
                            blended_b = color_b * opacity + (1 - opacity) * 1.0
                            final_color = (blended_r, blended_g, blended_b)
                        else:
                            final_color = (color_r, color_g, color_b)
                        
                        logger.info(f"添加水印 - 位置: {position}, 坐标: ({x:.1f}, {y:.1f}), 文字: '{watermark_text}', 角度: {rotation}°, 透明度: {opacity}, 最终颜色: RGB({final_color[0]:.2f},{final_color[1]:.2f},{final_color[2]:.2f})")
                        
                        # 估算文本尺寸（用于居中计算）
                        text_width = len(watermark_text) * fontsize * 0.6  # 中文字符宽度
                        text_height = fontsize
                        
                        # 将角度映射到最接近的标准角度（0, 90, 180, 270）
                        # 因为TextWriter在单点模式下无法正常工作，统一使用insert_text
                        if rotation < 45:
                            standard_rotation = 0
                        elif rotation < 135:
                            standard_rotation = 90
                        elif rotation < 225:
                            standard_rotation = 180
                        elif rotation < 315:
                            standard_rotation = 270
                        else:
                            standard_rotation = 0
                        
                        if rotation != standard_rotation:
                            logger.info(f"角度 {rotation}° 映射为标准角度 {standard_rotation}°")
                        
                        # 计算插入位置
                        # insert_text的基准点在文本左下角
                        if position == 'center':
                            # 居中：需要调整使文本中心在目标位置
                            if standard_rotation == 90 or standard_rotation == 270:
                                # 旋转90度后，宽度和高度互换
                                offset_x = -text_height / 2
                                offset_y = text_width / 2
                            else:
                                offset_x = -text_width / 2
                                offset_y = text_height / 2
                            insert_x = x + offset_x
                            insert_y = y - offset_y
                        else:
                            # 其他位置：直接使用目标位置
                            insert_x = x
                            insert_y = y
                        
                        logger.info(f"添加水印 - 位置: {position}, 目标: ({x:.1f}, {y:.1f}), 插入: ({insert_x:.1f}, {insert_y:.1f}), 角度: {standard_rotation}°")
                        
                        # 使用insert_text插入水印（可靠稳定）
                        page.insert_text(
                            (insert_x, insert_y),
                            watermark_text,
                            fontsize=fontsize,
                            fontname=fontname,
                            color=final_color,
                            render_mode=0,
                            rotate=standard_rotation
                        )
            
            # 保存
            output_filename = f"{timestamp}_output_{original_filename}"
            output_path = os.path.join(app.config['CONVERTED_FOLDER'], output_filename)
            
            if add_password:
                permissions = 0
                if allow_print:
                    permissions |= fitz.PDF_PERM_PRINT
                if allow_copy:
                    permissions |= fitz.PDF_PERM_COPY
                if allow_modify:
                    permissions |= fitz.PDF_PERM_MODIFY
                if allow_annotate:
                    permissions |= fitz.PDF_PERM_ANNOTATE
                
                doc.save(
                    output_path,
                    encryption=fitz.PDF_ENCRYPT_AES_256,
                    user_pw=user_pw,
                    owner_pw=owner_pw,
                    permissions=permissions,
                    garbage=4,
                    deflate=True
                )
            else:
                doc.save(output_path, garbage=4, deflate=True)
            
            output_size = os.path.getsize(output_path)
            output_size_mb = output_size / (1024 * 1024)
            elapsed = time.time() - start_time
            
            doc.close()
            os.remove(input_path)
            
            logger.info(f"处理完成: {output_filename}, 耗时{elapsed:.2f}秒")
            
            return jsonify({
                'success': True,
                'filename': output_filename,
                'total_pages': total_pages,
                'file_size': output_size,
                'file_size_mb': round(output_size_mb, 2),
                'elapsed_time': round(elapsed, 2)
            })
            
        except Exception as e:
            logger.error(f"处理失败: {str(e)}")
            logger.error(traceback.format_exc())
            # 确保文档被关闭
            try:
                if 'doc' in locals():
                    doc.close()
            except:
                pass
            # 删除临时文件
            if os.path.exists(input_path):
                try:
                    os.remove(input_path)
                except Exception as remove_error:
                    logger.warning(f"无法删除临时文件: {remove_error}")
            return jsonify({'error': f'处理失败: {str(e)}'}), 500
            
    except Exception as e:
        logger.error(f"错误: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': f'处理失败: {str(e)}'}), 500


@app.route('/pdf/remove-watermark', methods=['POST'])
def remove_watermark():
    """
    PDF去水印功能
    支持：文本水印删除、指定区域覆盖、图像水印删除
    """
    try:
        if 'file' not in request.files:
            return jsonify({'error': '请选择PDF文件'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': '请选择PDF文件'}), 400
        
        # 获取参数
        remove_mode = request.form.get('remove_mode', 'text')  # text/area/image/auto
        watermark_text = request.form.get('watermark_text', '')  # 要删除的水印文字
        
        # 区域模式参数
        area_x = float(request.form.get('area_x', '0'))
        area_y = float(request.form.get('area_y', '0'))
        area_width = float(request.form.get('area_width', '0'))
        area_height = float(request.form.get('area_height', '0'))
        
        # 颜色匹配参数（用于自动识别）
        match_color = request.form.get('match_color', 'false') == 'true'
        color_r = float(request.form.get('color_r', '0.8'))
        color_g = float(request.form.get('color_g', '0.8'))
        color_b = float(request.form.get('color_b', '0.8'))
        color_tolerance = float(request.form.get('color_tolerance', '0.1'))
        
        # 保存文件
        timestamp = int(time.time() * 1000)
        original_filename = secure_filename(file.filename)
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{timestamp}_{original_filename}")
        file.save(input_path)
        
        try:
            start_time = time.time()
            file_size = os.path.getsize(input_path)
            file_size_mb = file_size / (1024 * 1024)
            
            if file_size_mb > 50:
                os.remove(input_path)
                return jsonify({'error': '文件大小超过50MB限制'}), 400
            
            doc = fitz.open(input_path)
            total_pages = len(doc)
            
            if total_pages == 0:
                doc.close()
                os.remove(input_path)
                return jsonify({'error': 'PDF文件为空'}), 400
            
            if total_pages > 100:
                doc.close()
                os.remove(input_path)
                return jsonify({'error': '页数超过100页限制'}), 400
            
            logger.info(f"去水印: {original_filename}, 页数: {total_pages}, 模式: {remove_mode}")
            
            removed_count = 0
            
            for page_num in range(total_pages):
                page = doc[page_num]
                rect = page.rect
                
                if remove_mode == 'text' and watermark_text:
                    # 文本水印删除模式 - 基于颜色识别水印
                    text_dict = page.get_text("dict")
                    
                    watermark_rects = []
                    for block in text_dict.get("blocks", []):
                        if "lines" in block:
                            for line in block["lines"]:
                                for span in line["spans"]:
                                    text = span.get("text", "").strip()
                                    color = span.get("color", 0)
                                    
                                    # 转换颜色
                                    if isinstance(color, int):
                                        r = ((color >> 16) & 0xFF) / 255
                                        g = ((color >> 8) & 0xFF) / 255
                                        b = (color & 0xFF) / 255
                                    else:
                                        r, g, b = 0, 0, 0
                                    
                                    # 检查是否包含水印文字且是浅色（水印特征）
                                    is_light_color = (r > 0.6 and g > 0.6 and b > 0.6)
                                    contains_watermark = watermark_text in text
                                    
                                    if contains_watermark and is_light_color:
                                        bbox = span.get("bbox")
                                        if bbox:
                                            watermark_rects.append(fitz.Rect(bbox))
                                            logger.info(f"文本匹配 - 找到水印: '{text}' RGB({r:.2f},{g:.2f},{b:.2f})")
                    
                    logger.info(f"文本匹配 - 第{page_num+1}页找到 {len(watermark_rects)} 处水印")
                    
                    # 使用redact删除水印
                    for rect in watermark_rects:
                        page.add_redact_annot(rect)
                        removed_count += 1
                    
                    if watermark_rects:
                        page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)
                    
                elif remove_mode == 'area' and area_width > 0 and area_height > 0:
                    # 区域覆盖模式（按比例计算实际位置）
                    # 页面尺寸
                    page_w = rect.width
                    page_h = rect.height
                    
                    # 计算实际坐标
                    actual_x = page_w * area_x / 100
                    actual_y = page_h * area_y / 100
                    actual_w = page_w * area_width / 100
                    actual_h = page_h * area_height / 100
                    
                    logger.info(f"区域覆盖 - 页面尺寸: {page_w}x{page_h}")
                    logger.info(f"区域覆盖 - 输入参数: x={area_x}%, y={area_y}%, w={area_width}%, h={area_height}%")
                    logger.info(f"区域覆盖 - 实际坐标: x={actual_x:.1f}, y={actual_y:.1f}, w={actual_w:.1f}, h={actual_h:.1f}")
                    
                    # 创建覆盖区域
                    area_rect = fitz.Rect(
                        actual_x, actual_y,
                        actual_x + actual_w, actual_y + actual_h
                    )
                    
                    logger.info(f"区域覆盖 - Rect: {area_rect}")
                    
                    # 添加白色矩形覆盖（不使用redact，改用绘制白色矩形）
                    shape = page.new_shape()
                    shape.draw_rect(area_rect)
                    shape.finish(color=(1, 1, 1), fill=(1, 1, 1))
                    shape.commit()
                    removed_count += 1
                    
                elif remove_mode == 'auto':
                    # 自动识别模式：根据文本颜色特征识别水印（浅灰色）
                    text_dict = page.get_text("dict")
                    
                    watermark_rects = []
                    for block in text_dict.get("blocks", []):
                        if "lines" in block:
                            for line in block["lines"]:
                                for span in line["spans"]:
                                    span_color = span.get("color", 0)
                                    text = span.get("text", "").strip()
                                    
                                    # 转换颜色值
                                    if isinstance(span_color, int):
                                        r = ((span_color >> 16) & 0xFF) / 255
                                        g = ((span_color >> 8) & 0xFF) / 255
                                        b = (span_color & 0xFF) / 255
                                    else:
                                        r, g, b = 0, 0, 0
                                    
                                    # 判断是否为浅色（可能是水印）- 颜色接近灰色
                                    is_light_gray = (abs(r - 0.8) < 0.15 and 
                                                    abs(g - 0.8) < 0.15 and 
                                                    abs(b - 0.8) < 0.15)
                                    
                                    should_remove = False
                                    if match_color:
                                        # 匹配指定颜色
                                        color_match = (
                                            abs(r - color_r) < color_tolerance and
                                            abs(g - color_g) < color_tolerance and
                                            abs(b - color_b) < color_tolerance
                                        )
                                        should_remove = color_match
                                    elif is_light_gray and text:
                                        should_remove = True
                                    
                                    if should_remove:
                                        bbox = span.get("bbox")
                                        if bbox:
                                            watermark_rects.append(fitz.Rect(bbox))
                                            logger.info(f"自动识别 - 水印: '{text}' RGB({r:.2f},{g:.2f},{b:.2f})")
                                            removed_count += 1
                    
                    logger.info(f"自动识别 - 第{page_num+1}页找到 {len(watermark_rects)} 处水印")
                    
                    # 使用redact删除水印
                    for rect in watermark_rects:
                        page.add_redact_annot(rect)
                    
                    if watermark_rects:
                        page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)
                    
                elif remove_mode == 'image':
                    # 图像水印删除模式
                    image_list = page.get_images()
                    for img_index, img in enumerate(image_list):
                        xref = img[0]
                        # 获取图像位置
                        img_rects = page.get_image_rects(xref)
                        for img_rect in img_rects:
                            # 检查图像是否覆盖大部分页面（可能是水印）
                            img_area = img_rect.width * img_rect.height
                            page_area = rect.width * rect.height
                            coverage = img_area / page_area
                            
                            # 如果图像覆盖超过50%的页面，可能是背景水印
                            if coverage > 0.5:
                                page.add_redact_annot(img_rect, fill=(1, 1, 1))
                                removed_count += 1
                    
                    page.apply_redactions()
            
            # 保存结果
            output_filename = f"{timestamp}_no_watermark_{original_filename}"
            output_path = os.path.join(app.config['CONVERTED_FOLDER'], output_filename)
            doc.save(output_path, garbage=4, deflate=True)
            
            output_size = os.path.getsize(output_path)
            output_size_mb = output_size / (1024 * 1024)
            elapsed = time.time() - start_time
            
            doc.close()
            os.remove(input_path)
            
            logger.info(f"去水印完成: {output_filename}, 移除{removed_count}处, 耗时{elapsed:.2f}秒")
            
            return jsonify({
                'success': True,
                'filename': output_filename,
                'total_pages': total_pages,
                'removed_count': removed_count,
                'file_size': output_size,
                'file_size_mb': round(output_size_mb, 2),
                'elapsed_time': round(elapsed, 2)
            })
            
        except Exception as e:
            logger.error(f"去水印失败: {str(e)}")
            logger.error(traceback.format_exc())
            try:
                if 'doc' in locals():
                    doc.close()
            except:
                pass
            if os.path.exists(input_path):
                try:
                    os.remove(input_path)
                except:
                    pass
            return jsonify({'error': f'处理失败: {str(e)}'}), 500
            
    except Exception as e:
        logger.error(f"去水印错误: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': f'处理失败: {str(e)}'}), 500


@app.route('/pdf/remove-watermark-preview', methods=['POST'])
def remove_watermark_preview():
    """
    去水印预览 - 返回第一页的预览图
    """
    try:
        if 'file' not in request.files:
            return jsonify({'error': '请选择PDF文件'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': '请选择PDF文件'}), 400
        
        # 保存文件
        timestamp = int(time.time() * 1000)
        original_filename = secure_filename(file.filename)
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{timestamp}_preview_{original_filename}")
        file.save(input_path)
        
        try:
            doc = fitz.open(input_path)
            if len(doc) == 0:
                doc.close()
                os.remove(input_path)
                return jsonify({'error': 'PDF文件为空'}), 400
            
            # 渲染第一页
            page = doc[0]
            zoom = 1.5
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat)
            
            # 转为base64
            img_data = pix.tobytes("jpeg")
            img_base64 = base64.b64encode(img_data).decode()
            
            page_info = {
                'width': page.rect.width,
                'height': page.rect.height,
                'total_pages': len(doc)
            }
            
            doc.close()
            os.remove(input_path)
            
            return jsonify({
                'success': True,
                'preview': f"data:image/jpeg;base64,{img_base64}",
                'page_info': page_info
            })
            
        except Exception as e:
            logger.error(f"预览失败: {str(e)}")
            if os.path.exists(input_path):
                os.remove(input_path)
            return jsonify({'error': f'预览失败: {str(e)}'}), 500
            
    except Exception as e:
        logger.error(f"预览错误: {str(e)}")
        return jsonify({'error': f'处理失败: {str(e)}'}), 500


if __name__ == '__main__':
    logger.info("========================================")
    logger.info("PDF工具服务启动 v2.0")
    logger.info("端口: 8789")
    logger.info("支持的格式: PDF/Word/Excel/PPT/图片")
    logger.info("最大文件大小: 100MB")
    logger.info("========================================")
    logger.info("功能特性:")
    logger.info("  ✓ PDF转Word/Excel/PPT/图片")
    logger.info("  ✓ Word/Excel/PPT转PDF")
    logger.info("  ✓ PDF压缩优化")
    logger.info("  ✓ PDF页面编排")
    logger.info("  ✓ PDF旋转")
    logger.info("  ✓ 文件转长图")
    logger.info("  ✓ 长图打印")
    logger.info("  ✓ 转黑白/灰度")
    logger.info("  ✓ 文件重命名")
    logger.info("  ✓ 添加页码 (NEW!)")
    logger.info("  ✓ PDF页面裁剪 (NEW!)")
    logger.info("  ✓ PDF页面合并 (NEW!)")
    logger.info("  ✓ PDF加水印/密码 (NEW!)")
    logger.info("  ✓ PDF去水印 (NEW!)")
    logger.info("========================================")
    logger.info("API端点:")
    logger.info("  GET  /health - 健康检查")
    logger.info("  POST /pdf/add-page-numbers - 添加页码")
    logger.info("  POST /pdf/preview-page-number - 预览页码")
    logger.info("  POST /pdf/crop - PDF页面裁剪")
    logger.info("  POST /pdf/crop-preview - 裁剪预览")
    logger.info("  POST /pdf/combine-pages - PDF页面合并")
    logger.info("  POST /pdf/combine-preview - 合并预览")
    logger.info("  POST /pdf/watermark-password - PDF加水印/密码")
    logger.info("  POST /pdf/remove-watermark - PDF去水印")
    logger.info("  POST /file/rename - 文件重命名")
    logger.info("  POST /file/to-grayscale - 转黑白/灰度")
    logger.info("  POST /longimage/print - 长图打印")
    logger.info("  POST /pdf/compress - PDF压缩")
    logger.info("  POST /pdf/arrange/* - PDF页面编排")
    logger.info("  POST /file/to-long-image - 文件转长图")
    logger.info("  POST /pdf/rotate - PDF旋转")
    logger.info("  GET  /download/<filename> - 下载文件")
    logger.info("========================================\n")
    
    # 禁用reloader解决Windows权限问题
    app.run(host='0.0.0.0', port=8789, debug=True, use_reloader=False)
